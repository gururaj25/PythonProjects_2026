# generate_zip.py - FIXED VERSION
# Run: python generate_zip.py

import zipfile
from io import BytesIO


def create_zip():
    zip_buffer = BytesIO()
    files = {}

    # ================================================================
    # requirements.txt
    # ================================================================
    files["sql_consolidator/requirements.txt"] = (
        "sqlparse==0.4.4\n"
        "python-docx==1.1.0\n"
        "openpyxl==3.1.2\n"
        "streamlit==1.31.0\n"
        "pyyaml==6.0.1\n"
        "pandas==2.1.4\n"
        "colorama==0.4.6\n"
        "tqdm==4.66.1\n"
        "regex==2023.12.25\n"
        "chardet==5.2.0\n"
        "rich==13.7.0\n"
        "pytest==7.4.4\n"
        "pytest-cov==4.1.0\n"
    )

    # ================================================================
    # config/config.yaml
    # ================================================================
    files["sql_consolidator/config/config.yaml"] = (
        "# SQL Consolidator Configuration File\n"
        "\n"
        "scanner:\n"
        "  default_extensions:\n"
        "    - \".sql\"\n"
        "    - \".txt\"\n"
        "    - \".log\"\n"
        "    - \".bak\"\n"
        "  exclude_extensions:\n"
        "    - \".exe\"\n"
        "    - \".dll\"\n"
        "    - \".bin\"\n"
        "    - \".jpg\"\n"
        "    - \".png\"\n"
        "  max_file_size_mb: 100\n"
        "  encoding_fallbacks:\n"
        "    - \"utf-8\"\n"
        "    - \"latin-1\"\n"
        "    - \"cp1252\"\n"
        "    - \"ascii\"\n"
        "\n"
        "parser:\n"
        "  min_query_length: 10\n"
        "  max_query_length: 500000\n"
        "  supported_query_types:\n"
        "    - \"SELECT\"\n"
        "    - \"INSERT\"\n"
        "    - \"UPDATE\"\n"
        "    - \"DELETE\"\n"
        "    - \"MERGE\"\n"
        "    - \"CREATE\"\n"
        "    - \"ALTER\"\n"
        "    - \"DROP\"\n"
        "    - \"TRUNCATE\"\n"
        "    - \"EXEC\"\n"
        "    - \"EXECUTE\"\n"
        "    - \"WITH\"\n"
        "    - \"CALL\"\n"
        "  test_keywords:\n"
        "    - \"test_table\"\n"
        "    - \"tmp_test\"\n"
        "    - \"debug_log\"\n"
        "\n"
        "deduplicator:\n"
        "  similarity_threshold: 0.95\n"
        "  normalize_whitespace: true\n"
        "  normalize_case: true\n"
        "\n"
        "formatter:\n"
        "  keyword_case: \"upper\"\n"
        "  indent_width: 4\n"
        "  max_line_length: 120\n"
        "  reindent: true\n"
        "  strip_comments: false\n"
        "\n"
        "output:\n"
        "  generate_sql: true\n"
        "  generate_txt: true\n"
        "  generate_docx: true\n"
        "  generate_xlsx: true\n"
        "  generate_report: true\n"
        "  include_metadata: true\n"
        "\n"
        "analyzer:\n"
        "  enable_risk_detection: true\n"
        "  enable_complexity_scoring: true\n"
        "  enable_importance_ranking: true\n"
        "\n"
        "logging:\n"
        "  level: \"INFO\"\n"
        "  log_to_file: true\n"
        "  log_file: \"logs/sql_consolidator.log\"\n"
        "  max_log_size_mb: 10\n"
    )

    # ================================================================
    # config/__init__.py
    # ================================================================
    files["sql_consolidator/config/__init__.py"] = '"""Configuration Package"""\n'

    # ================================================================
    # scanner/__init__.py
    # ================================================================
    files["sql_consolidator/scanner/__init__.py"] = (
        '"""File Scanner Package"""\n'
        "from .file_scanner import FileScanner, ScannedFile, ScanResult\n"
        "__all__ = ['FileScanner', 'ScannedFile', 'ScanResult']\n"
    )

    # ================================================================
    # scanner/file_scanner.py
    # ================================================================
    files["sql_consolidator/scanner/file_scanner.py"] = '''"""
File Scanner Module
Recursively scans directories and extracts file content for SQL parsing.
"""

import os
import logging
import chardet
from pathlib import Path
from typing import List, Dict, Optional, Generator
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class ScannedFile:
    file_path: str
    file_name: str
    folder_path: str
    extension: str
    size_bytes: int
    content: str
    encoding: str
    line_count: int
    scan_status: str = "success"
    error_message: str = ""


@dataclass
class ScanResult:
    scanned_files: List[ScannedFile] = field(default_factory=list)
    failed_files: List[Dict] = field(default_factory=list)
    total_files_found: int = 0
    total_files_scanned: int = 0
    total_files_failed: int = 0
    total_size_bytes: int = 0


class FileScanner:
    def __init__(self, config: Dict):
        self.config = config
        self.scanner_config = config.get("scanner", {})
        self.default_extensions = set(
            self.scanner_config.get("default_extensions", [".sql", ".txt", ".log"])
        )
        self.exclude_extensions = set(
            self.scanner_config.get("exclude_extensions", [])
        )
        self.max_file_size = (
            self.scanner_config.get("max_file_size_mb", 100) * 1024 * 1024
        )
        self.encoding_fallbacks = self.scanner_config.get(
            "encoding_fallbacks", ["utf-8", "latin-1", "cp1252"]
        )

    def scan_directory(
        self,
        directory: str,
        extensions: Optional[List[str]] = None,
        keyword_filter: Optional[List[str]] = None,
        exclude_dirs: Optional[List[str]] = None,
    ) -> ScanResult:
        result = ScanResult()
        target_extensions = set(extensions) if extensions else self.default_extensions
        exclude_dirs = set(exclude_dirs or [".git", "__pycache__", "node_modules"])

        if not os.path.exists(directory):
            logger.error(f"Directory does not exist: {directory}")
            return result

        logger.info(f"Starting scan of directory: {directory}")
        all_files = list(self._walk_directory(directory, target_extensions, exclude_dirs))
        result.total_files_found = len(all_files)
        logger.info(f"Found {len(all_files)} files to process")

        for file_path in all_files:
            try:
                scanned = self._process_file(file_path)
                if scanned:
                    if keyword_filter:
                        if self._matches_keyword_filter(scanned.content, keyword_filter):
                            result.scanned_files.append(scanned)
                            result.total_size_bytes += scanned.size_bytes
                    else:
                        result.scanned_files.append(scanned)
                        result.total_size_bytes += scanned.size_bytes
                    result.total_files_scanned += 1
            except Exception as e:
                result.failed_files.append({"file_path": str(file_path), "error": str(e)})
                result.total_files_failed += 1
                logger.warning(f"Failed to process file {file_path}: {e}")

        logger.info(
            f"Scan complete. Scanned: {result.total_files_scanned}, "
            f"Failed: {result.total_files_failed}"
        )
        return result

    def _walk_directory(
        self, directory: str, extensions: set, exclude_dirs: set
    ) -> Generator[Path, None, None]:
        for root, dirs, files in os.walk(directory):
            dirs[:] = [d for d in dirs if d not in exclude_dirs and not d.startswith(".")]
            for file_name in files:
                file_path = Path(root) / file_name
                ext = file_path.suffix.lower()
                if ext in self.exclude_extensions:
                    continue
                if extensions and ext not in extensions:
                    continue
                try:
                    if file_path.stat().st_size > self.max_file_size:
                        logger.warning(f"Skipping large file: {file_path}")
                        continue
                except OSError:
                    continue
                yield file_path

    def _process_file(self, file_path: Path) -> Optional[ScannedFile]:
        file_stat = file_path.stat()
        content, encoding = self._read_file_with_encoding(file_path)
        if content is None:
            return None
        return ScannedFile(
            file_path=str(file_path.absolute()),
            file_name=file_path.name,
            folder_path=str(file_path.parent.absolute()),
            extension=file_path.suffix.lower(),
            size_bytes=file_stat.st_size,
            content=content,
            encoding=encoding,
            line_count=content.count("\\n") + 1,
        )

    def _read_file_with_encoding(self, file_path: Path):
        try:
            with open(file_path, "rb") as f:
                raw_data = f.read(min(32768, file_path.stat().st_size))
                detected = chardet.detect(raw_data)
                detected_encoding = detected.get("encoding", "utf-8")
        except Exception:
            detected_encoding = "utf-8"

        encodings_to_try = [detected_encoding] + self.encoding_fallbacks
        encodings_to_try = list(dict.fromkeys(e for e in encodings_to_try if e))

        for encoding in encodings_to_try:
            try:
                with open(file_path, "r", encoding=encoding, errors="strict") as f:
                    content = f.read()
                return content, encoding
            except (UnicodeDecodeError, LookupError):
                continue

        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            return content, "utf-8"
        except Exception as e:
            logger.error(f"Cannot read file {file_path}: {e}")
            return None, None

    @staticmethod
    def _matches_keyword_filter(content: str, keywords: List[str]) -> bool:
        content_lower = content.lower()
        return any(kw.lower() in content_lower for kw in keywords)

    def scan_multiple_directories(
        self,
        directories: List[str],
        extensions: Optional[List[str]] = None,
        keyword_filter: Optional[List[str]] = None,
    ) -> ScanResult:
        merged_result = ScanResult()
        for directory in directories:
            result = self.scan_directory(directory, extensions, keyword_filter)
            merged_result.scanned_files.extend(result.scanned_files)
            merged_result.failed_files.extend(result.failed_files)
            merged_result.total_files_found += result.total_files_found
            merged_result.total_files_scanned += result.total_files_scanned
            merged_result.total_files_failed += result.total_files_failed
            merged_result.total_size_bytes += result.total_size_bytes
        return merged_result
'''

    # ================================================================
    # parser/__init__.py
    # ================================================================
    files["sql_consolidator/parser/__init__.py"] = (
        '"""SQL Parser Package"""\n'
        "from .sql_parser import SQLParser, ExtractedQuery, QueryType\n"
        "__all__ = ['SQLParser', 'ExtractedQuery', 'QueryType']\n"
    )

    # ================================================================
    # parser/sql_parser.py
    # ================================================================
    files["sql_consolidator/parser/sql_parser.py"] = '''"""
SQL Parser Module
Extracts, classifies, and validates SQL queries from raw text content.
"""

import re
import logging
import hashlib
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum
import sqlparse

logger = logging.getLogger(__name__)


class QueryType(Enum):
    SELECT = "SELECT"
    INSERT = "INSERT"
    UPDATE = "UPDATE"
    DELETE = "DELETE"
    MERGE = "MERGE"
    CREATE = "CREATE"
    ALTER = "ALTER"
    DROP = "DROP"
    TRUNCATE = "TRUNCATE"
    EXEC = "EXEC"
    WITH = "WITH"
    TRANSACTION = "TRANSACTION"
    STORED_PROCEDURE = "STORED_PROCEDURE"
    VIEW = "VIEW"
    FUNCTION = "FUNCTION"
    TRIGGER = "TRIGGER"
    INDEX = "INDEX"
    UNKNOWN = "UNKNOWN"


@dataclass
class ExtractedQuery:
    raw_sql: str
    normalized_sql: str
    query_type: QueryType
    source_file: str
    source_folder: str
    line_number_start: int
    line_number_end: int
    table_names: List[str] = field(default_factory=list)
    has_joins: bool = False
    has_subquery: bool = False
    has_cte: bool = False
    has_aggregation: bool = False
    has_where_clause: bool = False
    is_complete: bool = True
    complexity_score: int = 0
    importance_score: int = 0
    risk_level: str = "LOW"
    is_duplicate: bool = False
    duplicate_of: str = ""
    fingerprint: str = ""


class LineTracker:
    def __init__(self, content: str):
        self.lines = content.split("\\n")

    def find_query_lines(self, query: str) -> Tuple[int, int]:
        first_line = query.split("\\n")[0].strip()[:50]
        for i, line in enumerate(self.lines, 1):
            if first_line and first_line in line:
                line_count = query.count("\\n") + 1
                return i, min(i + line_count - 1, len(self.lines))
        return 1, 1


class SQLParser:
    def __init__(self, config: Dict):
        self.config = config
        self.parser_config = config.get("parser", {})
        self.min_length = self.parser_config.get("min_query_length", 10)
        self.max_length = self.parser_config.get("max_query_length", 500000)
        self.test_keywords = self.parser_config.get("test_keywords", [])
        self.supported_types = self.parser_config.get(
            "supported_query_types",
            ["SELECT", "INSERT", "UPDATE", "DELETE", "CREATE", "ALTER", "DROP"]
        )

    def parse_file_content(
        self, content: str, file_path: str, folder_path: str
    ) -> List[ExtractedQuery]:
        extracted_queries = []
        try:
            sqlparse_queries = self._extract_with_sqlparse(content, file_path, folder_path)
            extracted_queries.extend(sqlparse_queries)
            if not sqlparse_queries:
                regex_queries = self._extract_with_regex(content, file_path, folder_path)
                extracted_queries.extend(regex_queries)
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {e}")

        valid_queries = []
        for query in extracted_queries:
            if self._is_valid_query(query):
                self._enrich_query_metadata(query)
                valid_queries.append(query)
        return valid_queries

    def _extract_with_sqlparse(
        self, content: str, file_path: str, folder_path: str
    ) -> List[ExtractedQuery]:
        queries = []
        line_tracker = LineTracker(content)
        try:
            statements = sqlparse.parse(content)
            for statement in statements:
                raw_sql = str(statement).strip()
                if not raw_sql or len(raw_sql) < self.min_length:
                    continue
                if self._is_comment_only(statement):
                    continue
                query_type = self._classify_query_type(raw_sql)
                if query_type == QueryType.UNKNOWN:
                    continue
                line_start, line_end = line_tracker.find_query_lines(raw_sql)
                queries.append(ExtractedQuery(
                    raw_sql=raw_sql,
                    normalized_sql=self._normalize_sql(raw_sql),
                    query_type=query_type,
                    source_file=file_path,
                    source_folder=folder_path,
                    line_number_start=line_start,
                    line_number_end=line_end,
                ))
        except Exception as e:
            logger.warning(f"sqlparse extraction failed for {file_path}: {e}")
        return queries

    def _extract_with_regex(
        self, content: str, file_path: str, folder_path: str
    ) -> List[ExtractedQuery]:
        queries = []
        lines = content.split("\\n")
        current_query_lines = []
        query_start_line = 0
        in_query = False

        sql_starters = re.compile(
            r"^\\s*(SELECT|INSERT|UPDATE|DELETE|MERGE|CREATE|ALTER|DROP|"
            r"TRUNCATE|EXEC|EXECUTE|WITH|BEGIN|DECLARE|CALL)\\b",
            re.IGNORECASE,
        )

        for line_num, line in enumerate(lines, 1):
            stripped = line.strip()
            if sql_starters.match(stripped) and not in_query:
                in_query = True
                query_start_line = line_num
                current_query_lines = [line]
            elif in_query:
                current_query_lines.append(line)
                if ";" in line or (
                    not stripped and len(current_query_lines) > 1
                    and any(ln.strip() for ln in current_query_lines[-3:])
                ):
                    raw_sql = "\\n".join(current_query_lines).strip()
                    if len(raw_sql) >= self.min_length:
                        query_type = self._classify_query_type(raw_sql)
                        if query_type != QueryType.UNKNOWN:
                            queries.append(ExtractedQuery(
                                raw_sql=raw_sql,
                                normalized_sql=self._normalize_sql(raw_sql),
                                query_type=query_type,
                                source_file=file_path,
                                source_folder=folder_path,
                                line_number_start=query_start_line,
                                line_number_end=line_num,
                            ))
                    in_query = False
                    current_query_lines = []

        if in_query and current_query_lines:
            raw_sql = "\\n".join(current_query_lines).strip()
            if len(raw_sql) >= self.min_length:
                query_type = self._classify_query_type(raw_sql)
                if query_type != QueryType.UNKNOWN:
                    queries.append(ExtractedQuery(
                        raw_sql=raw_sql,
                        normalized_sql=self._normalize_sql(raw_sql),
                        query_type=query_type,
                        source_file=file_path,
                        source_folder=folder_path,
                        line_number_start=query_start_line,
                        line_number_end=len(lines),
                    ))
        return queries

    def _classify_query_type(self, sql: str) -> QueryType:
        sql_upper = sql.upper().strip()
        sql_upper = re.sub(
            r"^(/\\*.*?\\*/\\s*|--[^\\n]*\\n\\s*)*", "", sql_upper, flags=re.DOTALL
        )
        type_patterns = [
            (r"^\\s*CREATE\\s+(OR\\s+REPLACE\\s+)?PROCEDURE\\b", QueryType.STORED_PROCEDURE),
            (r"^\\s*CREATE\\s+(OR\\s+REPLACE\\s+)?FUNCTION\\b", QueryType.FUNCTION),
            (r"^\\s*CREATE\\s+(OR\\s+REPLACE\\s+)?TRIGGER\\b", QueryType.TRIGGER),
            (r"^\\s*CREATE\\s+(OR\\s+REPLACE\\s+)?VIEW\\b", QueryType.VIEW),
            (r"^\\s*CREATE\\s+(OR\\s+REPLACE\\s+)?INDEX\\b", QueryType.INDEX),
            (r"^\\s*CREATE\\b", QueryType.CREATE),
            (r"^\\s*ALTER\\b", QueryType.ALTER),
            (r"^\\s*DROP\\b", QueryType.DROP),
            (r"^\\s*TRUNCATE\\b", QueryType.TRUNCATE),
            (r"^\\s*SELECT\\b", QueryType.SELECT),
            (r"^\\s*WITH\\s+\\w+.*\\bAS\\s*\\(", QueryType.WITH),
            (r"^\\s*INSERT\\b", QueryType.INSERT),
            (r"^\\s*UPDATE\\b", QueryType.UPDATE),
            (r"^\\s*DELETE\\b", QueryType.DELETE),
            (r"^\\s*MERGE\\b", QueryType.MERGE),
            (r"^\\s*EXEC(UTE)?\\b", QueryType.EXEC),
            (r"^\\s*(BEGIN|COMMIT|ROLLBACK)\\b", QueryType.TRANSACTION),
            (r"^\\s*DECLARE\\s+@", QueryType.EXEC),
            (r"^\\s*CALL\\b", QueryType.EXEC),
        ]
        for pattern, query_type in type_patterns:
            if re.match(pattern, sql_upper):
                return query_type
        return QueryType.UNKNOWN

    def _is_valid_query(self, query: ExtractedQuery) -> bool:
        sql = query.raw_sql.strip()
        if len(sql) < self.min_length or len(sql) > self.max_length:
            return False
        if not sql.replace("\\n", "").replace("\\t", "").strip():
            return False
        if self._is_pure_comment(sql):
            return False
        if not self._has_minimum_sql_structure(sql):
            return False
        return True

    def _has_minimum_sql_structure(self, sql: str) -> bool:
        sql_upper = sql.upper().strip()
        keywords = [
            "SELECT", "INSERT", "UPDATE", "DELETE", "CREATE", "ALTER",
            "DROP", "TRUNCATE", "MERGE", "EXEC", "WITH", "BEGIN", "CALL"
        ]
        return any(sql_upper.startswith(kw) for kw in keywords)

    def _is_comment_only(self, statement) -> bool:
        try:
            for token in statement.tokens:
                if token.ttype not in (
                    sqlparse.tokens.Comment.Single,
                    sqlparse.tokens.Comment.Multiline,
                    sqlparse.tokens.Newline,
                    sqlparse.tokens.Whitespace,
                ):
                    return False
            return True
        except Exception:
            return False

    def _is_pure_comment(self, sql: str) -> bool:
        lines = sql.strip().split("\\n")
        non_comment = [
            ln.strip() for ln in lines
            if ln.strip()
            and not ln.strip().startswith("--")
            and not ln.strip().startswith("/*")
            and not ln.strip().startswith("*")
            and not ln.strip().startswith("*/")
        ]
        return len(non_comment) == 0

    def _normalize_sql(self, sql: str) -> str:
        sql = re.sub(r"--[^\\n]*", "", sql)
        sql = re.sub(r"/\\*.*?\\*/", "", sql, flags=re.DOTALL)
        sql = " ".join(sql.split()).upper().rstrip(";").strip()
        return sql

    def _enrich_query_metadata(self, query: ExtractedQuery) -> None:
        sql = query.raw_sql.upper()
        query.has_joins = bool(re.search(
            r"\\b(INNER|LEFT|RIGHT|FULL|CROSS|OUTER)\\s+JOIN\\b|\\bJOIN\\b", sql
        ))
        query.has_subquery = bool(re.search(r"\\(\\s*SELECT\\b", sql))
        query.has_cte = query.query_type == QueryType.WITH or bool(
            re.search(r"\\bWITH\\s+\\w+\\s+AS\\s*\\(", sql)
        )
        query.has_aggregation = bool(re.search(
            r"\\b(COUNT|SUM|AVG|MAX|MIN|GROUP\\s+BY|HAVING)\\b", sql
        ))
        query.has_where_clause = bool(re.search(r"\\bWHERE\\b", sql))
        query.table_names = self._extract_table_names(query.raw_sql)
        query.fingerprint = self._generate_fingerprint(query.normalized_sql)
        query.complexity_score = self._calculate_complexity(query)
        query.risk_level = self._assess_risk_level(query)
        query.importance_score = self._calculate_importance(query)

    def _extract_table_names(self, sql: str) -> List[str]:
        table_names = []
        patterns = [
            r"\\bFROM\\s+([\\w\\.]+)",
            r"\\bJOIN\\s+([\\w\\.]+)",
            r"\\bINTO\\s+([\\w\\.]+)",
            r"\\bUPDATE\\s+([\\w\\.]+)",
            r"\\bTABLE\\s+([\\w\\.]+)",
        ]
        skip_words = {
            "SELECT", "SET", "WHERE", "GROUP", "ORDER", "HAVING",
            "ON", "AND", "OR", "NOT", "NULL", "AS", "IF", "EXISTS"
        }
        for pattern in patterns:
            matches = re.findall(pattern, sql, re.IGNORECASE)
            for match in matches:
                clean = re.sub(r"[`\\[\\]\\\"\\']", "", match).strip()
                if clean and clean.upper() not in skip_words:
                    table_names.append(clean)
        return list(set(table_names))

    def _generate_fingerprint(self, normalized_sql: str) -> str:
        return hashlib.sha256(normalized_sql.encode("utf-8")).hexdigest()

    def _calculate_complexity(self, query: ExtractedQuery) -> int:
        score = 1
        sql = query.raw_sql.upper()
        score += len(re.findall(r"\\bJOIN\\b", sql)) * 3
        score += len(re.findall(r"\\(\\s*SELECT\\b", sql)) * 4
        if query.has_cte:
            score += 3
        if query.has_aggregation:
            score += 2
        if re.search(r"\\bOVER\\s*\\(", sql):
            score += 4
        score += len(re.findall(r"\\bUNION\\b", sql)) * 2
        if re.search(r"\\bHAVING\\b", sql):
            score += 2
        score += min(len(query.raw_sql) // 200, 10)
        return score

    def _assess_risk_level(self, query: ExtractedQuery) -> str:
        if query.query_type in (QueryType.DELETE, QueryType.UPDATE):
            if not query.has_where_clause:
                return "CRITICAL"
        if query.query_type == QueryType.DROP:
            return "HIGH"
        if query.query_type == QueryType.TRUNCATE:
            return "HIGH"
        if query.query_type in (QueryType.DELETE, QueryType.UPDATE):
            return "MEDIUM"
        if query.query_type == QueryType.INSERT:
            return "MEDIUM"
        return "LOW"

    def _calculate_importance(self, query: ExtractedQuery) -> int:
        score = 0
        type_scores = {
            QueryType.CREATE: 8, QueryType.ALTER: 7,
            QueryType.STORED_PROCEDURE: 9, QueryType.VIEW: 8,
            QueryType.FUNCTION: 8, QueryType.TRIGGER: 7,
            QueryType.SELECT: 5, QueryType.INSERT: 6,
            QueryType.UPDATE: 6, QueryType.DELETE: 6,
            QueryType.MERGE: 7, QueryType.WITH: 6,
        }
        score += type_scores.get(query.query_type, 3)
        score += min(query.complexity_score // 3, 10)
        if query.has_joins:
            score += 3
        if query.has_subquery:
            score += 4
        if query.has_cte:
            score += 3
        if query.has_aggregation:
            score += 2
        if query.has_where_clause:
            score += 1
        score += min(len(query.table_names), 5)
        return score
'''

    # ================================================================
    # deduplicator/__init__.py
    # ================================================================
    files["sql_consolidator/deduplicator/__init__.py"] = (
        '"""Deduplicator Package"""\n'
        "from .deduplicator import QueryDeduplicator, DeduplicationResult\n"
        "__all__ = ['QueryDeduplicator', 'DeduplicationResult']\n"
    )

    # ================================================================
    # deduplicator/deduplicator.py
    # ================================================================
    files["sql_consolidator/deduplicator/deduplicator.py"] = '''"""
Deduplicator Module
Identifies and removes duplicate SQL queries using multiple strategies.
"""

import re
import logging
import hashlib
from typing import List, Dict, Tuple
from dataclasses import dataclass, field
from difflib import SequenceMatcher

logger = logging.getLogger(__name__)


@dataclass
class DeduplicationResult:
    unique_queries: List = field(default_factory=list)
    duplicate_groups: List[Dict] = field(default_factory=list)
    total_input: int = 0
    total_unique: int = 0
    total_duplicates: int = 0
    exact_duplicates: int = 0
    semantic_duplicates: int = 0


class QueryDeduplicator:
    def __init__(self, config: Dict):
        self.config = config
        self.dedup_config = config.get("deduplicator", {})
        self.similarity_threshold = self.dedup_config.get("similarity_threshold", 0.95)

    def deduplicate(self, queries: List) -> DeduplicationResult:
        result = DeduplicationResult()
        result.total_input = len(queries)
        if not queries:
            return result

        logger.info(f"Starting deduplication of {len(queries)} queries")

        after_exact, exact_groups = self._exact_deduplication(queries)
        result.exact_duplicates = result.total_input - len(after_exact)
        result.duplicate_groups.extend(exact_groups)

        after_norm, norm_groups = self._normalized_deduplication(after_exact)
        result.exact_duplicates += len(after_exact) - len(after_norm)
        result.duplicate_groups.extend(norm_groups)

        if len(after_norm) <= 5000:
            after_sem, sem_groups = self._semantic_deduplication(after_norm)
            result.semantic_duplicates = len(after_norm) - len(after_sem)
            result.duplicate_groups.extend(sem_groups)
            result.unique_queries = after_sem
        else:
            result.unique_queries = after_norm

        result.total_unique = len(result.unique_queries)
        result.total_duplicates = result.total_input - result.total_unique
        logger.info(
            f"Deduplication complete. {result.total_unique} unique "
            f"from {result.total_input} ({result.total_duplicates} removed)"
        )
        return result

    def _exact_deduplication(self, queries: List) -> Tuple[List, List[Dict]]:
        seen: Dict[str, object] = {}
        unique: List = []
        groups: List[Dict] = []
        for query in queries:
            fp = query.fingerprint
            if fp in seen:
                query.is_duplicate = True
                query.duplicate_of = seen[fp].source_file
                existing = next((g for g in groups if g["fingerprint"] == fp), None)
                if existing:
                    existing["duplicates"].append({
                        "file": query.source_file,
                        "line": query.line_number_start,
                        "type": "EXACT"
                    })
                else:
                    groups.append({
                        "fingerprint": fp,
                        "original": seen[fp].source_file,
                        "original_line": seen[fp].line_number_start,
                        "duplicates": [{"file": query.source_file,
                                        "line": query.line_number_start,
                                        "type": "EXACT"}],
                        "duplicate_type": "EXACT",
                        "query_preview": query.raw_sql[:100]
                    })
            else:
                seen[fp] = query
                unique.append(query)
        return unique, groups

    def _normalized_deduplication(self, queries: List) -> Tuple[List, List[Dict]]:
        seen: Dict[str, object] = {}
        unique: List = []
        groups: List[Dict] = []
        for query in queries:
            key = self._get_normalized_key(query.raw_sql)
            norm_fp = hashlib.sha256(key.encode()).hexdigest()
            if norm_fp in seen:
                query.is_duplicate = True
                query.duplicate_of = seen[norm_fp].source_file
                groups.append({
                    "fingerprint": norm_fp,
                    "original": seen[norm_fp].source_file,
                    "original_line": seen[norm_fp].line_number_start,
                    "duplicates": [{"file": query.source_file,
                                    "line": query.line_number_start,
                                    "type": "NORMALIZED"}],
                    "duplicate_type": "FORMATTING",
                    "query_preview": query.raw_sql[:100]
                })
            else:
                seen[norm_fp] = query
                unique.append(query)
        return unique, groups

    def _semantic_deduplication(self, queries: List) -> Tuple[List, List[Dict]]:
        type_groups: Dict[str, List] = {}
        for query in queries:
            k = query.query_type.value
            type_groups.setdefault(k, []).append(query)
        all_unique: List = []
        all_groups: List[Dict] = []
        for _, grp in type_groups.items():
            u, g = self._semantic_dedup_group(grp)
            all_unique.extend(u)
            all_groups.extend(g)
        return all_unique, all_groups

    def _semantic_dedup_group(self, queries: List) -> Tuple[List, List[Dict]]:
        unique: List = []
        groups: List[Dict] = []
        token_cache = [self._tokenize_query(q.raw_sql) for q in queries]
        for i, query in enumerate(queries):
            is_dup = False
            for j, uq in enumerate(unique):
                sim = self._calculate_similarity(
                    token_cache[i], self._tokenize_query(uq.raw_sql)
                )
                if sim >= self.similarity_threshold:
                    query.is_duplicate = True
                    query.duplicate_of = uq.source_file
                    is_dup = True
                    groups.append({
                        "fingerprint": f"sem_{i}_{j}",
                        "original": uq.source_file,
                        "original_line": uq.line_number_start,
                        "duplicates": [{"file": query.source_file,
                                        "line": query.line_number_start,
                                        "type": "SEMANTIC",
                                        "similarity": f"{sim:.2%}"}],
                        "duplicate_type": "SEMANTIC",
                        "query_preview": query.raw_sql[:100]
                    })
                    break
            if not is_dup:
                unique.append(query)
        return unique, groups

    def _get_normalized_key(self, sql: str) -> str:
        sql = re.sub(r"--[^\\n]*", " ", sql)
        sql = re.sub(r"/\\*.*?\\*/", " ", sql, flags=re.DOTALL)
        sql = sql.upper()
        sql = re.sub(r"\\s+", " ", sql).strip().rstrip(";").strip()
        sql = re.sub(r"\'[^\']*\'", "\'?\'", sql)
        sql = re.sub(r"\\b\\d+\\b", "?", sql)
        return sql

    def _tokenize_query(self, sql: str) -> List[str]:
        return re.findall(r"\\b\\w+\\b", self._get_normalized_key(sql))

    def _calculate_similarity(self, t1: List[str], t2: List[str]) -> float:
        if not t1 or not t2:
            return 0.0
        if min(len(t1), len(t2)) / max(len(t1), len(t2)) < 0.5:
            return 0.0
        return SequenceMatcher(None, " ".join(t1), " ".join(t2)).ratio()

    def generate_dedup_report(self, result: DeduplicationResult) -> str:
        lines = [
            "=" * 80,
            "DEDUPLICATION REPORT",
            "=" * 80,
            f"Total queries processed : {result.total_input}",
            f"Unique queries retained : {result.total_unique}",
            f"Total duplicates removed: {result.total_duplicates}",
            f"  - Exact duplicates    : {result.exact_duplicates}",
            f"  - Semantic duplicates : {result.semantic_duplicates}",
            f"Deduplication rate      : "
            f"{(result.total_duplicates / max(result.total_input, 1)) * 100:.1f}%",
            "", "DUPLICATE GROUPS:", "-" * 80,
        ]
        for i, group in enumerate(result.duplicate_groups[:50], 1):
            lines.append(f"\\nGroup #{i} ({group['duplicate_type']})")
            lines.append(f"  Original : {group['original']} (Line {group['original_line']})")
            lines.append(f"  Preview  : {group['query_preview'][:80]}")
            for dup in group.get("duplicates", []):
                lines.append(
                    f"  Duplicate: {dup['file']} (Line {dup['line']}) [{dup['type']}]"
                )
        return "\\n".join(lines)
'''

    # ================================================================
    # formatter/__init__.py
    # ================================================================
    files["sql_consolidator/formatter/__init__.py"] = (
        '"""SQL Formatter Package"""\n'
        "from .sql_formatter import SQLFormatter\n"
        "__all__ = ['SQLFormatter']\n"
    )

    # ================================================================
    # formatter/sql_formatter.py
    # ================================================================
    files["sql_consolidator/formatter/sql_formatter.py"] = '''"""
SQL Formatter Module
Professional SQL formatting with proper indentation and structure.
"""

import re
import logging
from typing import List, Dict
from datetime import datetime
import sqlparse

logger = logging.getLogger(__name__)


class SQLFormatter:
    def __init__(self, config: Dict):
        self.config = config
        self.fmt_config = config.get("formatter", {})
        self.keyword_case = self.fmt_config.get("keyword_case", "upper")
        self.indent_width = self.fmt_config.get("indent_width", 4)
        self.reindent = self.fmt_config.get("reindent", True)
        self.strip_comments = self.fmt_config.get("strip_comments", False)

    def format_query(self, query) -> str:
        try:
            formatted = sqlparse.format(
                query.raw_sql,
                reindent=self.reindent,
                keyword_case=self.keyword_case,
                identifier_case="lower",
                strip_comments=self.strip_comments,
                indent_width=self.indent_width,
                indent_tabs=False,
                use_space_around_operators=True,
            )
            return self._post_process(formatted)
        except Exception as e:
            logger.warning(f"Formatting error: {e}")
            return query.raw_sql

    def _post_process(self, sql: str) -> str:
        sql = sql.rstrip()
        if not sql.endswith(";"):
            sql += ";"
        sql = re.sub(r",(?!\\s)", ", ", sql)
        sql = re.sub(r"\\n{3,}", "\\n\\n", sql)
        return sql

    def format_all_queries(self, queries: List) -> Dict[str, List[str]]:
        categorized: Dict[str, List[str]] = {}
        for query in queries:
            category = self._get_category_name(query.query_type)
            formatted = self.format_query(query)
            formatted_with_meta = self._add_metadata_comment(query, formatted)
            categorized.setdefault(category, []).append(formatted_with_meta)
        return categorized

    def _get_category_name(self, query_type) -> str:
        from sql_parser.sql_parser import QueryType
        mapping = {
            QueryType.SELECT: "SELECT Queries",
            QueryType.WITH: "CTE Queries",
            QueryType.INSERT: "INSERT Queries",
            QueryType.UPDATE: "UPDATE Queries",
            QueryType.DELETE: "DELETE Queries",
            QueryType.MERGE: "MERGE Queries",
            QueryType.CREATE: "DDL - CREATE",
            QueryType.ALTER: "DDL - ALTER",
            QueryType.DROP: "DDL - DROP",
            QueryType.TRUNCATE: "DDL - TRUNCATE",
            QueryType.EXEC: "EXEC",
            QueryType.STORED_PROCEDURE: "STORED PROCEDURES",
            QueryType.VIEW: "VIEWS",
            QueryType.FUNCTION: "FUNCTIONS",
            QueryType.TRIGGER: "TRIGGERS",
            QueryType.INDEX: "INDEXES",
            QueryType.TRANSACTION: "TRANSACTION BLOCKS",
        }
        return mapping.get(query_type, "Uncategorized")

    def _add_metadata_comment(self, query, formatted_sql: str) -> str:
        meta = [
            f"-- Source File : {query.source_file}",
            f"-- Query Type  : {query.query_type.value}",
            f"-- Line Number : {query.line_number_start} - {query.line_number_end}",
            f"-- Complexity  : {query.complexity_score}",
            f"-- Risk Level  : {query.risk_level}",
            f"-- Importance  : {query.importance_score}",
        ]
        if query.table_names:
            meta.append(f"-- Tables      : {', '.join(query.table_names)}")
        sep = "-" * 80
        return f"\\n{sep}\\n" + "\\n".join(meta) + f"\\n{sep}\\n{formatted_sql}\\n"

    def generate_master_sql_document(self, queries: List, scan_metadata: Dict) -> str:
        sections = [self._generate_header(scan_metadata)]
        categorized = self.format_all_queries(queries)
        order = [
            "DDL - CREATE", "DDL - ALTER", "DDL - DROP", "DDL - TRUNCATE",
            "VIEWS", "FUNCTIONS", "STORED PROCEDURES", "EXEC", "TRIGGERS",
            "INDEXES", "SELECT Queries", "CTE Queries", "INSERT Queries",
            "UPDATE Queries", "DELETE Queries", "MERGE Queries",
            "TRANSACTION BLOCKS", "Uncategorized",
        ]
        for section_name in order:
            if section_name in categorized:
                sections.append(self._generate_section(section_name, categorized[section_name]))
        for cat, qlist in categorized.items():
            if cat not in order:
                sections.append(self._generate_section(cat, qlist))
        sections.append(self._generate_footer())
        return "\\n\\n".join(sections)

    def _generate_header(self, metadata: Dict) -> str:
        lines = [
            "/*", "=" * 76,
            "  SQL QUERY CONSOLIDATION MASTER DOCUMENT",
            "  Generated by SQL Consolidator Tool v1.0",
            "=" * 76, "",
            f"  Generated On      : {metadata.get('scan_datetime', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}",
            f"  Input Directory   : {metadata.get('input_directory', 'N/A')}",
            f"  Files Scanned     : {metadata.get('total_files_scanned', 0)}",
            f"  Queries Found     : {metadata.get('total_queries_found', 0)}",
            f"  Unique Queries    : {metadata.get('total_unique_queries', 0)}",
            f"  Duplicates Removed: {metadata.get('total_duplicates_removed', 0)}",
            "", "  QUERY DISTRIBUTION:",
        ]
        for qtype, count in metadata.get("query_type_distribution", {}).items():
            lines.append(f"    {qtype:<30} : {count}")
        lines += ["", "=" * 76, "*/", ""]
        return "\\n".join(lines)

    def _generate_section(self, name: str, queries: List[str]) -> str:
        header = "\\n".join([
            "", "/" + "=" * 79,
            f"||  {name.upper():<74} ||",
            f"||  Total: {len(queries):<70} ||",
            "=" * 79 + "/",
        ])
        return header + "\\n" + "\\nGO\\n".join(queries)

    def _generate_footer(self) -> str:
        return "\\n".join([
            "", "/*", "=" * 76,
            "  END OF MASTER SQL DOCUMENT",
            f"  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "=" * 76, "*/",
        ])
'''

    # ================================================================
    # analyzer/__init__.py
    # ================================================================
    files["sql_consolidator/analyzer/__init__.py"] = (
        '"""Query Analyzer Package"""\n'
        "from .query_analyzer import QueryAnalyzer, AnalysisReport\n"
        "__all__ = ['QueryAnalyzer', 'AnalysisReport']\n"
    )

    # ================================================================
    # analyzer/query_analyzer.py
    # ================================================================
    files["sql_consolidator/analyzer/query_analyzer.py"] = '''"""
Query Analyzer Module
Analyzes SQL queries for risk detection, complexity, and insights.
"""

import re
import logging
from typing import List, Dict, Tuple
from collections import Counter
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class AnalysisReport:
    total_queries: int = 0
    risk_summary: Dict[str, int] = field(default_factory=dict)
    high_risk_queries: List = field(default_factory=list)
    complex_queries: List = field(default_factory=list)
    most_common_tables: List[Tuple[str, int]] = field(default_factory=list)
    query_type_distribution: Dict[str, int] = field(default_factory=dict)
    average_complexity: float = 0.0
    natural_language_summary: str = ""
    recommendations: List[str] = field(default_factory=list)
    top_important_queries: List = field(default_factory=list)


class QueryAnalyzer:
    def __init__(self, config: Dict):
        self.config = config
        self.analyzer_config = config.get("analyzer", {})

    def analyze(self, queries: List) -> AnalysisReport:
        report = AnalysisReport()
        report.total_queries = len(queries)
        if not queries:
            return report

        logger.info(f"Analyzing {len(queries)} queries")
        report.query_type_distribution = dict(
            Counter(q.query_type.value for q in queries).most_common()
        )
        report.risk_summary = dict(Counter(q.risk_level for q in queries))
        report.high_risk_queries = sorted(
            [q for q in queries if q.risk_level in ("HIGH", "CRITICAL")],
            key=lambda q: 0 if q.risk_level == "CRITICAL" else 1
        )
        report.complex_queries = sorted(
            queries, key=lambda q: q.complexity_score, reverse=True
        )[:20]
        report.average_complexity = sum(q.complexity_score for q in queries) / len(queries)
        report.top_important_queries = sorted(
            queries, key=lambda q: q.importance_score, reverse=True
        )[:50]
        all_tables = []
        for q in queries:
            all_tables.extend(q.table_names)
        report.most_common_tables = Counter(
            t.lower() for t in all_tables if t
        ).most_common(20)
        report.recommendations = self._generate_recommendations(queries, report)
        report.natural_language_summary = self._generate_nl_summary(report)
        logger.info("Analysis complete")
        return report

    def _generate_recommendations(self, queries: List, report: AnalysisReport) -> List[str]:
        recs = []
        critical = report.risk_summary.get("CRITICAL", 0)
        if critical > 0:
            recs.append(
                f"CRITICAL: {critical} DELETE/UPDATE queries found WITHOUT WHERE clause. "
                f"Review immediately to prevent mass data loss."
            )
        high = report.risk_summary.get("HIGH", 0)
        if high > 0:
            recs.append(
                f"HIGH RISK: {high} DROP/TRUNCATE queries found. "
                f"Ensure these are intentional and have backups."
            )
        complex_q = [q for q in queries if q.complexity_score > 20]
        if complex_q:
            recs.append(
                f"{len(complex_q)} highly complex queries detected (score > 20). "
                f"Consider optimization review."
            )
        from sql_parser.sql_parser import QueryType
        risky_dml = [
            q for q in queries
            if q.query_type in (QueryType.UPDATE, QueryType.DELETE)
            and not q.has_where_clause
        ]
        if risky_dml:
            recs.append(
                f"{len(risky_dml)} UPDATE/DELETE queries lack WHERE clauses - "
                f"potential full table modifications!"
            )
        if report.most_common_tables:
            top = report.most_common_tables[0]
            recs.append(
                f"Most referenced table: '{top[0]}' appears in {top[1]} queries. "
                f"Ensure proper indexing."
            )
        ddl_count = sum(
            report.query_type_distribution.get(t, 0) for t in ["CREATE", "ALTER", "DROP"]
        )
        if ddl_count > 0:
            recs.append(
                f"{ddl_count} DDL statements found. Ensure schema changes are "
                f"version-controlled."
            )
        return recs

    def _generate_nl_summary(self, report: AnalysisReport) -> str:
        lines = [
            "NATURAL LANGUAGE SUMMARY", "=" * 60,
            f"This consolidated document contains {report.total_queries} unique SQL queries.",
            "",
        ]
        if report.query_type_distribution:
            lines.append("Query Breakdown:")
            for qtype, count in report.query_type_distribution.items():
                lines.append(f"  * {count} {qtype} queries")
            lines.append("")
        critical = report.risk_summary.get("CRITICAL", 0)
        high = report.risk_summary.get("HIGH", 0)
        medium = report.risk_summary.get("MEDIUM", 0)
        low = report.risk_summary.get("LOW", 0)
        lines += [
            "Risk Profile:",
            f"  * {critical} CRITICAL risk (DELETE/UPDATE without WHERE)",
            f"  * {high} HIGH risk (DROP/TRUNCATE)",
            f"  * {medium} MEDIUM risk (DML with WHERE)",
            f"  * {low} LOW risk (SELECT/DDL)",
            "",
            f"Average complexity score: {report.average_complexity:.1f}",
            "",
        ]
        if report.most_common_tables:
            lines.append("Top Referenced Tables:")
            for table, count in report.most_common_tables[:5]:
                lines.append(f"  * {table}: referenced {count} times")
            lines.append("")
        if report.recommendations:
            lines.append("Key Recommendations:")
            for rec in report.recommendations:
                lines.append(f"  {rec}")
        return "\\n".join(lines)
'''

    # ================================================================
    # reports/__init__.py
    # ================================================================
    files["sql_consolidator/reports/__init__.py"] = (
        '"""Report Generator Package"""\n'
        "from .report_generator import ReportGenerator\n"
        "__all__ = ['ReportGenerator']\n"
    )

    # ================================================================
    # reports/report_generator.py
    # ================================================================
    files["sql_consolidator/reports/report_generator.py"] = '''"""
Report Generator Module
Generates comprehensive reports in SQL, TXT, DOCX, and XLSX formats.
"""

import logging
from datetime import datetime
from typing import List, Dict
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

from formatter.sql_formatter import SQLFormatter

logger = logging.getLogger(__name__)


class ReportGenerator:
    def __init__(self, config: Dict, output_dir: str):
        self.config = config
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.formatter = SQLFormatter(config)
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    def generate_all_reports(
        self, queries: List, dedup_result, analysis_report, scan_metadata: Dict
    ) -> Dict[str, str]:
        output_files = {}
        cfg = self.config.get("output", {})

        if cfg.get("generate_sql", True):
            path = self._gen_master_sql(queries, scan_metadata)
            output_files["master_sql"] = path
            logger.info(f"Master SQL: {path}")

        if cfg.get("generate_txt", True):
            path = self._gen_text_report(queries, dedup_result, analysis_report, scan_metadata)
            output_files["text_report"] = path
            logger.info(f"Text report: {path}")

        if cfg.get("generate_docx", True):
            path = self._gen_word_doc(queries, dedup_result, analysis_report, scan_metadata)
            output_files["word_document"] = path
            logger.info(f"Word document: {path}")

        if cfg.get("generate_xlsx", True):
            path = self._gen_excel(queries, dedup_result, analysis_report, scan_metadata)
            output_files["excel_report"] = path
            logger.info(f"Excel report: {path}")

        output_files["duplicate_report"] = self._gen_dup_report(dedup_result)
        output_files["error_log"] = self._gen_error_log(scan_metadata)
        return output_files

    def _gen_master_sql(self, queries: List, metadata: Dict) -> str:
        path = str(self.output_dir / f"master_consolidated_{self.timestamp}.sql")
        content = self.formatter.generate_master_sql_document(queries, metadata)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return path

    def _gen_text_report(
        self, queries: List, dedup, analysis, metadata: Dict
    ) -> str:
        path = str(self.output_dir / f"summary_report_{self.timestamp}.txt")
        lines = [
            "=" * 80, "SQL QUERY CONSOLIDATION - SUMMARY REPORT", "=" * 80,
            f"Generated: {metadata.get('scan_datetime', 'N/A')}",
            "",
            "SCAN STATISTICS", "-" * 40,
            f"Input Directory    : {metadata.get('input_directory', 'N/A')}",
            f"Files Found        : {metadata.get('total_files_found', 0)}",
            f"Files Scanned      : {metadata.get('total_files_scanned', 0)}",
            f"Files Failed       : {metadata.get('total_files_failed', 0)}",
            "",
            "QUERY STATISTICS", "-" * 40,
            f"Total Queries Found    : {metadata.get('total_queries_found', 0)}",
            f"Unique Queries         : {dedup.total_unique}",
            f"Duplicates Removed     : {dedup.total_duplicates}",
            f"  Exact Duplicates     : {dedup.exact_duplicates}",
            f"  Semantic Duplicates  : {dedup.semantic_duplicates}",
            f"Deduplication Rate     : "
            f"{(dedup.total_duplicates / max(dedup.total_input, 1)) * 100:.1f}%",
            "",
            "QUERY TYPE DISTRIBUTION", "-" * 40,
        ]
        for qtype, count in analysis.query_type_distribution.items():
            lines.append(f"  {qtype:<30}: {count}")
        lines += ["", "RISK ANALYSIS", "-" * 40]
        for risk_level, count in analysis.risk_summary.items():
            lines.append(f"  {risk_level:<20}: {count}")
        lines += ["", "TOP REFERENCED TABLES", "-" * 40]
        for table, count in analysis.most_common_tables[:15]:
            lines.append(f"  {table:<40}: {count} references")
        lines += ["", analysis.natural_language_summary, "", "RECOMMENDATIONS", "-" * 40]
        for rec in analysis.recommendations:
            lines.append(f"  {rec}")
        if analysis.high_risk_queries:
            lines += ["", "HIGH RISK QUERIES", "-" * 40]
            for q in analysis.high_risk_queries[:10]:
                lines += [
                    f"  Risk: {q.risk_level} | Type: {q.query_type.value}",
                    f"  File: {q.source_file} (Line {q.line_number_start})",
                    f"  SQL : {q.raw_sql[:150]}",
                    "",
                ]
        lines.append("=" * 80)
        with open(path, "w", encoding="utf-8") as f:
            f.write("\\n".join(lines))
        return path

    def _gen_word_doc(
        self, queries: List, dedup, analysis, metadata: Dict
    ) -> str:
        path = str(self.output_dir / f"sql_report_{self.timestamp}.docx")
        doc = Document()
        doc.core_properties.title = "SQL Query Consolidation Report"

        doc.add_heading("SQL Query Consolidation Report", 0)
        doc.add_paragraph(f"Generated: {metadata.get('scan_datetime', 'N/A')}")
        doc.add_paragraph(f"Input: {metadata.get('input_directory', 'N/A')}")

        doc.add_page_break()
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(
            f"Files scanned: {metadata.get('total_files_scanned', 0)} | "
            f"Queries found: {metadata.get('total_queries_found', 0)} | "
            f"Unique queries: {dedup.total_unique} | "
            f"Duplicates removed: {dedup.total_duplicates}"
        )

        doc.add_heading("Statistics", level=2)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.rows[0].cells[0].text = "Metric"
        tbl.rows[0].cells[1].text = "Value"
        stats = [
            ("Files Scanned", str(metadata.get("total_files_scanned", 0))),
            ("Total Queries Found", str(metadata.get("total_queries_found", 0))),
            ("Unique Queries", str(dedup.total_unique)),
            ("Duplicates Removed", str(dedup.total_duplicates)),
            ("Critical Risk", str(analysis.risk_summary.get("CRITICAL", 0))),
            ("High Risk", str(analysis.risk_summary.get("HIGH", 0))),
            ("Avg Complexity", f"{analysis.average_complexity:.1f}"),
        ]
        for metric, value in stats:
            row = tbl.add_row()
            row.cells[0].text = metric
            row.cells[1].text = value

        doc.add_page_break()
        doc.add_heading("Risk Analysis", level=1)
        if analysis.high_risk_queries:
            for q in analysis.high_risk_queries[:10]:
                doc.add_paragraph(
                    f"[{q.risk_level}] {q.query_type.value} | "
                    f"{Path(q.source_file).name} (Line {q.line_number_start}): "
                    f"{q.raw_sql[:100]}...",
                    style="List Bullet"
                )
        else:
            doc.add_paragraph("No high-risk queries detected.")

        doc.add_heading("Recommendations", level=1)
        for rec in analysis.recommendations:
            doc.add_paragraph(rec, style="List Bullet")

        doc.add_page_break()
        doc.add_heading("Consolidated SQL Queries", level=1)
        categorized = self.formatter.format_all_queries(queries)
        for category, qlist in sorted(categorized.items()):
            doc.add_heading(f"{category} ({len(qlist)})", level=2)
            for i, qsql in enumerate(qlist[:50], 1):
                p = doc.add_paragraph()
                run = p.add_run(f"-- Query {i}\\n{qsql[:800]}")
                run.font.name = "Courier New"
                run.font.size = Pt(8)

        doc.save(path)
        return path

    def _gen_excel(
        self, queries: List, dedup, analysis, metadata: Dict
    ) -> str:
        path = str(self.output_dir / f"sql_analysis_{self.timestamp}.xlsx")
        wb = Workbook()

        # Summary sheet
        ws = wb.active
        ws.title = "Summary"
        ws["A1"] = "SQL QUERY CONSOLIDATION REPORT"
        ws["A1"].font = Font(bold=True, size=14, color="1F4E79")
        ws.merge_cells("A1:C1")
        ws["A2"] = f"Generated: {metadata.get('scan_datetime', 'N/A')}"
        ws.merge_cells("A2:C2")

        stats_rows = [
            ("", ""), ("SCAN STATISTICS", ""),
            ("Files Found", metadata.get("total_files_found", 0)),
            ("Files Scanned", metadata.get("total_files_scanned", 0)),
            ("Files Failed", metadata.get("total_files_failed", 0)),
            ("", ""), ("QUERY STATISTICS", ""),
            ("Total Queries Found", metadata.get("total_queries_found", 0)),
            ("Unique Queries", dedup.total_unique),
            ("Duplicates Removed", dedup.total_duplicates),
            ("Exact Duplicates", dedup.exact_duplicates),
            ("Semantic Duplicates", dedup.semantic_duplicates),
            ("", ""), ("RISK SUMMARY", ""),
            ("Critical Risk", analysis.risk_summary.get("CRITICAL", 0)),
            ("High Risk", analysis.risk_summary.get("HIGH", 0)),
            ("Medium Risk", analysis.risk_summary.get("MEDIUM", 0)),
            ("Low Risk", analysis.risk_summary.get("LOW", 0)),
            ("", ""), ("QUERY DISTRIBUTION", ""),
        ]
        for qtype, count in analysis.query_type_distribution.items():
            stats_rows.append((qtype, count))

        section_hdrs = {"SCAN STATISTICS", "QUERY STATISTICS", "RISK SUMMARY", "QUERY DISTRIBUTION"}
        row = 4
        for label, value in stats_rows:
            ws.cell(row=row, column=1, value=label)
            ws.cell(row=row, column=2, value=value)
            if label in section_hdrs:
                c = ws.cell(row=row, column=1)
                c.font = Font(bold=True, color="FFFFFF")
                c.fill = PatternFill(fill_type="solid", fgColor="2E75B6")
                ws.merge_cells(f"A{row}:B{row}")
            row += 1
        ws.column_dimensions["A"].width = 30
        ws.column_dimensions["B"].width = 20

        # All Queries sheet
        ws2 = wb.create_sheet("All Queries")
        headers = [
            "Query #", "Query Type", "Source File", "Folder",
            "Line Start", "Line End", "Tables", "Has JOINs",
            "Has Subquery", "Has CTE", "Complexity", "Importance",
            "Risk Level", "SQL Preview"
        ]
        hf = Font(bold=True, color="FFFFFF")
        hfill = PatternFill(fill_type="solid", fgColor="1F4E79")
        for col, hdr in enumerate(headers, 1):
            c = ws2.cell(row=1, column=col, value=hdr)
            c.font = hf
            c.fill = hfill
            c.alignment = Alignment(horizontal="center")
        ws2.freeze_panes = "A2"
        risk_colors = {
            "CRITICAL": "FF0000", "HIGH": "FFC000",
            "MEDIUM": "FFFF00", "LOW": "92D050"
        }
        for i, q in enumerate(queries, 1):
            r = i + 1
            vals = [
                i, q.query_type.value, q.source_file, q.source_folder,
                q.line_number_start, q.line_number_end,
                ", ".join(q.table_names[:5]),
                "Yes" if q.has_joins else "No",
                "Yes" if q.has_subquery else "No",
                "Yes" if q.has_cte else "No",
                q.complexity_score, q.importance_score, q.risk_level,
                q.raw_sql[:200].replace("\\n", " "),
            ]
            for col, val in enumerate(vals, 1):
                cell = ws2.cell(row=r, column=col, value=val)
                if col == 13:
                    cell.fill = PatternFill(
                        fill_type="solid",
                        fgColor=risk_colors.get(str(val), "FFFFFF")
                    )
        ws2.auto_filter.ref = f"A1:{get_column_letter(len(headers))}1"
        for col, w in enumerate([8,20,50,40,10,10,40,10,12,10,12,12,12,60], 1):
            ws2.column_dimensions[get_column_letter(col)].width = w

        # High Risk sheet
        ws3 = wb.create_sheet("High Risk Queries")
        rh = ["Risk Level", "Query Type", "Source File", "Line", "SQL Preview", "Issue"]
        rf = PatternFill(fill_type="solid", fgColor="C00000")
        for col, h in enumerate(rh, 1):
            c = ws3.cell(row=1, column=col, value=h)
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = rf
        for i, q in enumerate(analysis.high_risk_queries, 2):
            issue = (
                "DELETE/UPDATE without WHERE - may affect ALL rows!"
                if q.risk_level == "CRITICAL"
                else "DROP/TRUNCATE - irreversible operation"
            )
            ws3.cell(row=i, column=1, value=q.risk_level)
            ws3.cell(row=i, column=2, value=q.query_type.value)
            ws3.cell(row=i, column=3, value=q.source_file)
            ws3.cell(row=i, column=4, value=q.line_number_start)
            ws3.cell(row=i, column=5, value=q.raw_sql[:200].replace("\\n", " "))
            ws3.cell(row=i, column=6, value=issue)
        for col, w in enumerate([15, 20, 50, 10, 60, 50], 1):
            ws3.column_dimensions[get_column_letter(col)].width = w

        # Table Analysis sheet
        ws4 = wb.create_sheet("Table Analysis")
        th = ["Rank", "Table Name", "Reference Count"]
        tf = PatternFill(fill_type="solid", fgColor="375623")
        for col, h in enumerate(th, 1):
            c = ws4.cell(row=1, column=col, value=h)
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = tf
        for i, (tbl, cnt) in enumerate(analysis.most_common_tables, 2):
            ws4.cell(row=i, column=1, value=i - 1)
            ws4.cell(row=i, column=2, value=tbl)
            ws4.cell(row=i, column=3, value=cnt)

        # Duplicates sheet
        ws5 = wb.create_sheet("Duplicates")
        dh = ["Group #", "Dup Type", "Original File", "Orig Line",
              "Dup File", "Dup Line", "Match", "Preview"]
        df = PatternFill(fill_type="solid", fgColor="7030A0")
        for col, h in enumerate(dh, 1):
            c = ws5.cell(row=1, column=col, value=h)
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = df
        row = 2
        for i, grp in enumerate(dedup.duplicate_groups[:1000], 1):
            for dup in grp.get("duplicates", []):
                ws5.cell(row=row, column=1, value=i)
                ws5.cell(row=row, column=2, value=grp.get("duplicate_type", ""))
                ws5.cell(row=row, column=3, value=grp.get("original", ""))
                ws5.cell(row=row, column=4, value=grp.get("original_line", ""))
                ws5.cell(row=row, column=5, value=dup.get("file", ""))
                ws5.cell(row=row, column=6, value=dup.get("line", ""))
                ws5.cell(row=row, column=7, value=dup.get("type", ""))
                ws5.cell(row=row, column=8, value=grp.get("query_preview", "")[:150])
                row += 1

        wb.save(path)
        return path

    def _gen_dup_report(self, dedup) -> str:
        path = str(self.output_dir / f"duplicate_report_{self.timestamp}.txt")
        from deduplicator.deduplicator import QueryDeduplicator
        report = QueryDeduplicator(self.config).generate_dedup_report(dedup)
        with open(path, "w", encoding="utf-8") as f:
            f.write(report)
        return path

    def _gen_error_log(self, metadata: Dict) -> str:
        path = str(self.output_dir / f"error_log_{self.timestamp}.txt")
        lines = [f"ERROR LOG - {metadata.get('scan_datetime', 'N/A')}", "=" * 60, ""]
        failed = metadata.get("failed_files", [])
        if failed:
            lines.append(f"Failed Files ({len(failed)}):")
            for fi in failed:
                lines.append(f"  File : {fi.get('file_path', 'N/A')}")
                lines.append(f"  Error: {fi.get('error', 'Unknown')}")
                lines.append("")
        else:
            lines.append("No errors encountered.")
        with open(path, "w", encoding="utf-8") as f:
            f.write("\\n".join(lines))
        return path
'''

    # ================================================================
    # ui/__init__.py
    # ================================================================
    files["sql_consolidator/ui/__init__.py"] = '"""UI Package"""\n'

    # ================================================================
    # ui/streamlit_app.py
    # ================================================================
    files["sql_consolidator/ui/streamlit_app.py"] = '''"""
Streamlit GUI for SQL Query Consolidation Tool
Run with: streamlit run ui/streamlit_app.py
"""

import os
import sys
from pathlib import Path
import streamlit as st
import pandas as pd

sys.path.insert(0, str(Path(__file__).parent.parent))

from main import load_config, run_consolidation


def main():
    st.set_page_config(
        page_title="SQL Consolidator",
        page_icon="\\U0001f5c4",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    config = load_config("config/config.yaml")

    # Sidebar
    st.sidebar.title("Configuration")
    st.sidebar.markdown("---")

    input_dir = st.sidebar.text_input("Input Directory", placeholder="/path/to/sql/files")
    output_dir = st.sidebar.text_input("Output Directory", value="./output")

    extensions = st.sidebar.multiselect(
        "Include Extensions",
        options=[".sql", ".txt", ".log", ".bak", ".csv", ".json"],
        default=[".sql", ".txt", ".log"],
    )

    keyword_input = st.sidebar.text_area(
        "Keyword Filter (one per line)", placeholder="production\\ncritical"
    )
    keywords = (
        [k.strip() for k in keyword_input.split("\\n") if k.strip()]
        if keyword_input.strip() else None
    )

    skip_dedup = st.sidebar.checkbox("Skip Deduplication", value=False)
    min_complexity = st.sidebar.slider("Min Complexity Score", 0, 50, 0)
    min_importance = st.sidebar.slider("Min Importance Score", 0, 30, 0)

    st.sidebar.markdown("---")
    gen_sql = st.sidebar.checkbox("Generate .sql", value=True)
    gen_txt = st.sidebar.checkbox("Generate .txt", value=True)
    gen_docx = st.sidebar.checkbox("Generate .docx", value=True)
    gen_xlsx = st.sidebar.checkbox("Generate .xlsx", value=True)

    # Main content
    st.title("SQL Query Consolidation and Deduplication Tool")
    st.markdown(
        "Scan multiple files and directories to extract, deduplicate, "
        "and consolidate SQL queries into a single clean master document."
    )

    col1, col2, col3, col4 = st.columns(4)
    col1.info("Multi-file Scanning")
    col2.info("Smart Deduplication")
    col3.info("SQL Formatting")
    col4.info("Risk Analysis")

    st.markdown("---")

    if st.button("Start SQL Consolidation", type="primary", use_container_width=True):
        if not input_dir:
            st.error("Please specify an input directory")
            return
        if not os.path.exists(input_dir):
            st.error(f"Input directory not found: {input_dir}")
            return

        config["output"] = {
            "generate_sql": gen_sql,
            "generate_txt": gen_txt,
            "generate_docx": gen_docx,
            "generate_xlsx": gen_xlsx,
        }

        with st.spinner("Processing SQL files..."):
            try:
                results = run_consolidation(
                    input_dir=input_dir,
                    output_dir=output_dir,
                    config=config,
                    extensions=extensions,
                    keywords=keywords,
                    skip_dedup=skip_dedup,
                    min_complexity=min_complexity,
                    min_importance=min_importance,
                )
                if results:
                    st.success("Consolidation completed successfully!")
                    st.session_state["results"] = results
                else:
                    st.warning("No queries found in the specified directory")
            except Exception as e:
                st.error(f"Error during processing: {str(e)}")
                st.exception(e)

    if "results" in st.session_state:
        results = st.session_state["results"]
        dedup = results["dedup_result"]
        analysis = results["analysis_report"]

        st.markdown("---")
        st.subheader("Results Overview")
        c1, c2, c3, c4, c5 = st.columns(5)
        c1.metric("Files Scanned", results["scan_result"].total_files_scanned)
        c2.metric("Queries Found", len(results["all_queries"]))
        c3.metric("Unique Queries", dedup.total_unique,
                  delta=f"-{dedup.total_duplicates} dupes")
        c4.metric("Critical Risks", analysis.risk_summary.get("CRITICAL", 0))
        c5.metric("Time", f"{results.get('elapsed_seconds', 0):.1f}s")

        tab1, tab2, tab3, tab4 = st.tabs([
            "Analytics", "Query Browser", "Risk Analysis", "Downloads"
        ])

        with tab1:
            if analysis.query_type_distribution:
                st.subheader("Query Type Distribution")
                df = pd.DataFrame(
                    list(analysis.query_type_distribution.items()),
                    columns=["Query Type", "Count"]
                ).sort_values("Count", ascending=False)
                st.bar_chart(df.set_index("Query Type"))

            if analysis.most_common_tables:
                st.subheader("Most Referenced Tables")
                tdf = pd.DataFrame(
                    analysis.most_common_tables, columns=["Table", "References"]
                )
                st.bar_chart(tdf.set_index("Table"))

        with tab2:
            st.subheader("Query Browser")
            queries = results["unique_queries"]
            type_filter = st.selectbox(
                "Filter by Type",
                ["All"] + list(set(q.query_type.value for q in queries))
            )
            risk_filter = st.selectbox(
                "Filter by Risk", ["All", "CRITICAL", "HIGH", "MEDIUM", "LOW"]
            )
            filtered = [
                q for q in queries
                if (type_filter == "All" or q.query_type.value == type_filter)
                and (risk_filter == "All" or q.risk_level == risk_filter)
            ]
            st.caption(f"Showing {len(filtered)} of {len(queries)} queries")
            if filtered:
                data = [{
                    "Type": q.query_type.value,
                    "Risk": q.risk_level,
                    "Complexity": q.complexity_score,
                    "Importance": q.importance_score,
                    "Tables": ", ".join(q.table_names[:3]),
                    "File": Path(q.source_file).name,
                    "SQL Preview": q.raw_sql[:120].replace("\\n", " "),
                } for q in filtered[:500]]
                st.dataframe(pd.DataFrame(data), use_container_width=True, height=350)

                idx = st.selectbox(
                    "Select Query to View",
                    range(min(len(filtered), 500)),
                    format_func=lambda i: (
                        f"[{filtered[i].query_type.value}] "
                        f"{filtered[i].raw_sql[:60]}..."
                    )
                )
                ca, cb = st.columns([2, 1])
                with ca:
                    st.code(filtered[idx].raw_sql, language="sql")
                with cb:
                    q = filtered[idx]
                    st.json({
                        "Type": q.query_type.value,
                        "Risk": q.risk_level,
                        "Complexity": q.complexity_score,
                        "JOINs": q.has_joins,
                        "Subquery": q.has_subquery,
                        "CTE": q.has_cte,
                        "WHERE": q.has_where_clause,
                        "Tables": q.table_names,
                    })

        with tab3:
            st.subheader("Risk Analysis")
            r1, r2, r3, r4 = st.columns(4)
            r1.metric("CRITICAL", analysis.risk_summary.get("CRITICAL", 0))
            r2.metric("HIGH", analysis.risk_summary.get("HIGH", 0))
            r3.metric("MEDIUM", analysis.risk_summary.get("MEDIUM", 0))
            r4.metric("LOW", analysis.risk_summary.get("LOW", 0))

            if analysis.high_risk_queries:
                st.warning(f"Found {len(analysis.high_risk_queries)} HIGH/CRITICAL queries!")
                with st.expander("View High Risk Queries"):
                    rdata = [{
                        "Risk Level": q.risk_level,
                        "Query Type": q.query_type.value,
                        "File": Path(q.source_file).name,
                        "Line": q.line_number_start,
                        "SQL": q.raw_sql[:200].replace("\\n", " "),
                    } for q in analysis.high_risk_queries[:20]]
                    st.dataframe(pd.DataFrame(rdata), use_container_width=True)

            if analysis.recommendations:
                st.subheader("Recommendations")
                for rec in analysis.recommendations:
                    if "CRITICAL" in rec:
                        st.error(rec)
                    elif "HIGH" in rec or "risk" in rec.lower():
                        st.warning(rec)
                    else:
                        st.info(rec)

        with tab4:
            st.subheader("Download Reports")
            output_files = results.get("output_files", {})
            mime_map = {
                ".sql": "text/plain",
                ".txt": "text/plain",
                ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            }
            cols = st.columns(min(len(output_files), 3))
            for i, (rtype, fpath) in enumerate(output_files.items()):
                with cols[i % 3]:
                    if os.path.exists(fpath):
                        with open(fpath, "rb") as f:
                            data = f.read()
                        ext = Path(fpath).suffix
                        st.download_button(
                            label=f"Download {rtype.replace('_', ' ').title()}",
                            data=data,
                            file_name=Path(fpath).name,
                            mime=mime_map.get(ext, "application/octet-stream"),
                            use_container_width=True,
                        )


if __name__ == "__main__":
    main()
'''

    # ================================================================
    # main.py
    # ================================================================
    files["sql_consolidator/main.py"] = '''"""
SQL Query Consolidation and Deduplication Tool
Main entry point for command-line execution.
"""

import os
import sys
import logging
import argparse
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict
import yaml
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TimeElapsedColumn, BarColumn
from rich.panel import Panel
from rich.table import Table

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from scanner.file_scanner import FileScanner
from sql_parser.sql_parser import SQLParser, ExtractedQuery
from deduplicator.deduplicator import QueryDeduplicator, DeduplicationResult
from analyzer.query_analyzer import QueryAnalyzer
from reports.report_generator import ReportGenerator

console = Console()


def setup_logging(config: Dict) -> logging.Logger:
    log_config = config.get("logging", {})
    log_level = getattr(logging, log_config.get("level", "INFO"))
    log_file = log_config.get("log_file", "logs/sql_consolidator.log")
    Path(log_file).parent.mkdir(parents=True, exist_ok=True)
    handlers = [logging.StreamHandler()]
    if log_config.get("log_to_file", True):
        handlers.append(logging.FileHandler(log_file, encoding="utf-8"))
    logging.basicConfig(
        level=log_level,
        format="%(asctime)s | %(name)-20s | %(levelname)-8s | %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
        handlers=handlers,
    )
    return logging.getLogger("sql_consolidator")


def load_config(config_path: str = "config/config.yaml") -> Dict:
    if os.path.exists(config_path):
        with open(config_path, "r", encoding="utf-8") as f:
            return yaml.safe_load(f)
    logging.warning(f"Config not found: {config_path}. Using defaults.")
    return {}


def parse_arguments() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="SQL Query Consolidation and Deduplication Tool",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\\n"
            "  python main.py -i ./input -o ./output\\n"
            "  python main.py -i ./input -o ./output -e .sql .txt\\n"
            "  python main.py -i ./input -o ./output -k production\\n"
            "  python main.py -i ./input -o ./output --no-dedup\\n"
        ),
    )
    parser.add_argument("-i", "--input", required=True, help="Input directory path")
    parser.add_argument("-o", "--output", required=True, help="Output directory path")
    parser.add_argument("-e", "--extensions", nargs="+", default=None)
    parser.add_argument("-k", "--keywords", nargs="+", default=None)
    parser.add_argument("--exclude-dirs", nargs="+", default=None)
    parser.add_argument("--config", default="config/config.yaml")
    parser.add_argument("--no-dedup", action="store_true")
    parser.add_argument("--verbose", action="store_true")
    parser.add_argument("--min-complexity", type=int, default=0)
    parser.add_argument("--min-importance", type=int, default=0)
    return parser.parse_args()


def run_consolidation(
    input_dir: str,
    output_dir: str,
    config: Dict,
    extensions: Optional[List[str]] = None,
    keywords: Optional[List[str]] = None,
    exclude_dirs: Optional[List[str]] = None,
    skip_dedup: bool = False,
    min_complexity: int = 0,
    min_importance: int = 0,
) -> Dict:
    logger = logging.getLogger("sql_consolidator.pipeline")
    start_time = datetime.now()

    with Progress(
        SpinnerColumn(),
        "[progress.description]{task.description}",
        BarColumn(),
        TimeElapsedColumn(),
        console=console,
    ) as progress:

        # Step 1: Scan
        t1 = progress.add_task("[cyan]Scanning files...", total=None)
        scanner = FileScanner(config)
        scan_result = scanner.scan_directory(
            directory=input_dir,
            extensions=extensions,
            keyword_filter=keywords,
            exclude_dirs=exclude_dirs,
        )
        progress.update(t1, completed=True, description="[green]Files scanned")

        if scan_result.total_files_scanned == 0:
            console.print("[red]No files found. Check input directory.[/red]")
            return {}

        # Step 2: Parse
        t2 = progress.add_task(
            f"[cyan]Parsing {scan_result.total_files_scanned} files...",
            total=scan_result.total_files_scanned
        )
        parser = SQLParser(config)
        all_queries: List[ExtractedQuery] = []
        for sf in scan_result.scanned_files:
            extracted = parser.parse_file_content(sf.content, sf.file_path, sf.folder_path)
            all_queries.extend(extracted)
            progress.advance(t2)
        progress.update(t2, description=f"[green]{len(all_queries)} queries extracted")

        if min_complexity > 0 or min_importance > 0:
            orig = len(all_queries)
            all_queries = [
                q for q in all_queries
                if q.complexity_score >= min_complexity and q.importance_score >= min_importance
            ]
            logger.info(f"Filtered to {len(all_queries)} (removed {orig - len(all_queries)})")

        # Step 3: Dedup
        dedup_result = DeduplicationResult()
        if not skip_dedup:
            t3 = progress.add_task("[cyan]Deduplicating...", total=None)
            dedup_result = QueryDeduplicator(config).deduplicate(all_queries)
            unique_queries = dedup_result.unique_queries
            progress.update(t3, completed=True,
                            description=f"[green]{dedup_result.total_unique} unique")
        else:
            unique_queries = all_queries
            dedup_result.unique_queries = all_queries
            dedup_result.total_input = len(all_queries)
            dedup_result.total_unique = len(all_queries)

        # Step 4: Analyze
        t4 = progress.add_task("[cyan]Analyzing queries...", total=None)
        analysis_report = QueryAnalyzer(config).analyze(unique_queries)
        progress.update(t4, completed=True, description="[green]Analysis complete")

        # Step 5: Reports
        t5 = progress.add_task("[cyan]Generating reports...", total=None)
        scan_metadata = {
            "scan_datetime": start_time.strftime("%Y-%m-%d %H:%M:%S"),
            "input_directory": input_dir,
            "total_files_found": scan_result.total_files_found,
            "total_files_scanned": scan_result.total_files_scanned,
            "total_files_failed": scan_result.total_files_failed,
            "total_size_bytes": scan_result.total_size_bytes,
            "total_queries_found": len(all_queries),
            "total_unique_queries": dedup_result.total_unique,
            "total_duplicates_removed": dedup_result.total_duplicates,
            "total_errors": scan_result.total_files_failed,
            "failed_files": scan_result.failed_files,
            "query_type_distribution": analysis_report.query_type_distribution,
        }
        output_files = ReportGenerator(config, output_dir).generate_all_reports(
            queries=unique_queries,
            dedup_result=dedup_result,
            analysis_report=analysis_report,
            scan_metadata=scan_metadata,
        )
        progress.update(t5, completed=True, description="[green]Reports generated")

    elapsed = (datetime.now() - start_time).total_seconds()
    console.print(f"\\n[bold green]Complete in {elapsed:.1f} seconds[/bold green]")

    return {
        "scan_result": scan_result,
        "all_queries": all_queries,
        "unique_queries": unique_queries,
        "dedup_result": dedup_result,
        "analysis_report": analysis_report,
        "output_files": output_files,
        "elapsed_seconds": elapsed,
    }


def main():
    console.print(Panel(
        "  SQL QUERY CONSOLIDATION AND DEDUPLICATION TOOL\\n  Professional Edition v1.0",
        style="bold blue"
    ))
    args = parse_arguments()
    config = load_config(args.config)
    logger = setup_logging(config)

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    console.print(f"[bold]Input :[/bold] {args.input}")
    console.print(f"[bold]Output:[/bold] {args.output}")

    try:
        results = run_consolidation(
            input_dir=args.input,
            output_dir=args.output,
            config=config,
            extensions=args.extensions,
            keywords=args.keywords,
            exclude_dirs=args.exclude_dirs,
            skip_dedup=args.no_dedup,
            min_complexity=args.min_complexity,
            min_importance=args.min_importance,
        )
        if results:
            tbl = Table(title="Final Summary", header_style="bold green")
            tbl.add_column("Metric", style="cyan", width=35)
            tbl.add_column("Value", style="yellow", width=20)
            d = results["dedup_result"]
            a = results["analysis_report"]
            tbl.add_row("Total Queries Found", str(d.total_input))
            tbl.add_row("Unique Queries", str(d.total_unique))
            tbl.add_row("Duplicates Removed", str(d.total_duplicates))
            tbl.add_row(
                "Deduplication Rate",
                f"{(d.total_duplicates / max(d.total_input, 1)) * 100:.1f}%"
            )
            tbl.add_row("Avg Complexity", f"{a.average_complexity:.1f}")
            tbl.add_row("Critical Risk", str(a.risk_summary.get("CRITICAL", 0)))
            tbl.add_row("High Risk", str(a.risk_summary.get("HIGH", 0)))
            console.print(tbl)

            ftbl = Table(title="Generated Files", header_style="bold blue")
            ftbl.add_column("Report Type", style="cyan", width=25)
            ftbl.add_column("File Path", style="green", width=60)
            for rtype, fpath in results["output_files"].items():
                ftbl.add_row(rtype.replace("_", " ").title(), fpath)
            console.print(ftbl)

            if a.recommendations:
                console.print("\\n[bold yellow]Key Recommendations:[/bold yellow]")
                for rec in a.recommendations:
                    console.print(f"  {rec}")

    except KeyboardInterrupt:
        console.print("\\n[red]Interrupted.[/red]")
        sys.exit(1)
    except Exception as e:
        logger.error(f"Fatal error: {e}", exc_info=True)
        console.print(f"\\n[red]Error: {e}[/red]")
        sys.exit(1)


if __name__ == "__main__":
    main()
'''

    # ================================================================
    # tests/__init__.py
    # ================================================================
    files["sql_consolidator/tests/__init__.py"] = '"""Test Package"""\n'

    # ================================================================
    # tests/test_parser.py
    # ================================================================
    files["sql_consolidator/tests/test_parser.py"] = '''"""Unit tests for SQL Parser module."""
import pytest
import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from sql_parser.sql_parser import SQLParser, QueryType, ExtractedQuery


@pytest.fixture
def config():
    return {"parser": {"min_query_length": 10, "max_query_length": 500000,
                        "supported_query_types": ["SELECT","INSERT","UPDATE",
                        "DELETE","CREATE","ALTER","DROP","TRUNCATE","EXEC","WITH"],
                        "test_keywords": ["test_table","tmp_test"]}}

@pytest.fixture
def parser(config):
    return SQLParser(config)

def make_q(sql, qtype=None, p=None):
    p = p or SQLParser({"parser": {}})
    n = p._normalize_sql(sql)
    fp = p._generate_fingerprint(n)
    qt = qtype or p._classify_query_type(sql)
    return ExtractedQuery(raw_sql=sql, normalized_sql=n, query_type=qt,
                          source_file="t.sql", source_folder="/t",
                          line_number_start=1, line_number_end=1, fingerprint=fp)

class TestClassification:
    def test_select(self, parser): assert parser._classify_query_type("SELECT * FROM users") == QueryType.SELECT
    def test_insert(self, parser): assert parser._classify_query_type("INSERT INTO users (name) VALUES (\'x\')") == QueryType.INSERT
    def test_update(self, parser): assert parser._classify_query_type("UPDATE users SET name=\'x\' WHERE id=1") == QueryType.UPDATE
    def test_delete(self, parser): assert parser._classify_query_type("DELETE FROM users WHERE id=1") == QueryType.DELETE
    def test_create(self, parser): assert parser._classify_query_type("CREATE TABLE t (id INT)") == QueryType.CREATE
    def test_procedure(self, parser): assert parser._classify_query_type("CREATE PROCEDURE p AS SELECT 1") == QueryType.STORED_PROCEDURE
    def test_view(self, parser): assert parser._classify_query_type("CREATE VIEW v AS SELECT 1") == QueryType.VIEW
    def test_alter(self, parser): assert parser._classify_query_type("ALTER TABLE t ADD col INT") == QueryType.ALTER
    def test_drop(self, parser): assert parser._classify_query_type("DROP TABLE t") == QueryType.DROP
    def test_cte(self, parser): assert parser._classify_query_type("WITH cte AS (SELECT 1) SELECT * FROM cte") == QueryType.WITH
    def test_exec(self, parser): assert parser._classify_query_type("EXEC sp_test") == QueryType.EXEC
    def test_unknown(self, parser): assert parser._classify_query_type("GARBAGE TEXT") == QueryType.UNKNOWN

class TestValidation:
    def test_valid(self, parser):
        q = make_q("SELECT id FROM users WHERE active=1", p=parser)
        assert parser._is_valid_query(q) is True
    def test_empty(self, parser):
        q = make_q("SELECT * FROM users", p=parser)
        q.raw_sql = ""
        assert parser._is_valid_query(q) is False
    def test_comment_only(self, parser):
        assert parser._is_pure_comment("-- comment") is True
        assert parser._is_pure_comment("SELECT * FROM t") is False

class TestNormalization:
    def test_whitespace(self, parser):
        assert parser._normalize_sql("SELECT   *   FROM   t") == parser._normalize_sql("SELECT * FROM t")
    def test_case(self, parser):
        assert parser._normalize_sql("select * from t") == parser._normalize_sql("SELECT * FROM T")
    def test_semicolon(self, parser):
        assert parser._normalize_sql("SELECT * FROM t;") == parser._normalize_sql("SELECT * FROM t")

class TestMetadata:
    def test_joins(self, parser):
        q = make_q("SELECT u.id FROM users u INNER JOIN orders o ON u.id=o.uid", QueryType.SELECT, parser)
        parser._enrich_query_metadata(q)
        assert q.has_joins is True
    def test_subquery(self, parser):
        q = make_q("SELECT * FROM users WHERE id IN (SELECT uid FROM orders)", QueryType.SELECT, parser)
        parser._enrich_query_metadata(q)
        assert q.has_subquery is True
    def test_critical_risk(self, parser):
        q = make_q("DELETE FROM users", QueryType.DELETE, parser)
        q.has_where_clause = False
        assert parser._assess_risk_level(q) == "CRITICAL"
    def test_medium_risk(self, parser):
        q = make_q("UPDATE users SET x=1 WHERE id=1", QueryType.UPDATE, parser)
        q.has_where_clause = True
        assert parser._assess_risk_level(q) == "MEDIUM"
'''

    # ================================================================
    # tests/test_deduplicator.py
    # ================================================================
    files["sql_consolidator/tests/test_deduplicator.py"] = '''"""Unit tests for Deduplicator module."""
import pytest
import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from deduplicator.deduplicator import QueryDeduplicator
from sql_parser.sql_parser import SQLParser, ExtractedQuery, QueryType


@pytest.fixture
def config():
    return {"deduplicator": {"similarity_threshold": 0.95,
                              "normalize_whitespace": True, "normalize_case": True}}

@pytest.fixture
def deduplicator(config):
    return QueryDeduplicator(config)

def make_query(sql, source="test.sql", qtype=QueryType.SELECT):
    p = SQLParser({"parser": {}})
    n = p._normalize_sql(sql)
    fp = p._generate_fingerprint(n)
    return ExtractedQuery(raw_sql=sql, normalized_sql=n, query_type=qtype,
                          source_file=source, source_folder="/t",
                          line_number_start=1, line_number_end=1, fingerprint=fp)

class TestExact:
    def test_removes_exact_dup(self, deduplicator):
        q1, q2 = make_query("SELECT * FROM users"), make_query("SELECT * FROM users")
        u, g = deduplicator._exact_deduplication([q1, q2])
        assert len(u) == 1 and len(g) == 1

    def test_keeps_different(self, deduplicator):
        q1, q2 = make_query("SELECT * FROM users"), make_query("SELECT * FROM orders")
        u, g = deduplicator._exact_deduplication([q1, q2])
        assert len(u) == 2 and len(g) == 0

    def test_triple_dup(self, deduplicator):
        q1 = make_query("SELECT id FROM users")
        u, _ = deduplicator._exact_deduplication([q1, make_query("SELECT id FROM users"), make_query("SELECT id FROM users")])
        assert len(u) == 1

class TestNormalized:
    def test_case_insensitive(self, deduplicator):
        r = deduplicator.deduplicate([make_query("SELECT * FROM users"), make_query("select * from users")])
        assert r.total_unique == 1

    def test_whitespace_insensitive(self, deduplicator):
        r = deduplicator.deduplicate([make_query("SELECT * FROM users"), make_query("SELECT   *   FROM   users")])
        assert r.total_unique == 1

    def test_different_not_deduped(self, deduplicator):
        r = deduplicator.deduplicate([make_query("SELECT id FROM users WHERE active=1"), make_query("SELECT name FROM customers")])
        assert r.total_unique == 2

class TestStats:
    def test_statistics(self, deduplicator):
        queries = [make_query("SELECT * FROM users")] * 5 + [make_query("SELECT * FROM orders"), make_query("SELECT * FROM products")]
        r = deduplicator.deduplicate(queries)
        assert r.total_input == 7 and r.total_unique == 3 and r.total_duplicates == 4

    def test_empty_input(self, deduplicator):
        r = deduplicator.deduplicate([])
        assert r.total_input == 0 and r.total_unique == 0

    def test_single_query(self, deduplicator):
        r = deduplicator.deduplicate([make_query("SELECT * FROM users")])
        assert r.total_unique == 1 and r.total_duplicates == 0

    def test_report_content(self, deduplicator):
        r = deduplicator.deduplicate([make_query("SELECT * FROM users"), make_query("SELECT * FROM users"), make_query("INSERT INTO t VALUES (1)")])
        report = deduplicator.generate_dedup_report(r)
        assert "DEDUPLICATION REPORT" in report
'''

    # ================================================================
    # tests/test_scanner.py
    # ================================================================
    files["sql_consolidator/tests/test_scanner.py"] = '''"""Unit tests for File Scanner module."""
import os, pytest, tempfile
from pathlib import Path
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from scanner.file_scanner import FileScanner


@pytest.fixture
def config():
    return {"scanner": {"default_extensions": [".sql",".txt"],
                         "exclude_extensions": [".exe"],
                         "max_file_size_mb": 10,
                         "encoding_fallbacks": ["utf-8","latin-1"]}}

@pytest.fixture
def scanner(config):
    return FileScanner(config)

@pytest.fixture
def tmpdir_with_files():
    with tempfile.TemporaryDirectory() as d:
        (Path(d) / "test.sql").write_text("SELECT * FROM users;\\nINSERT INTO logs VALUES (1);", encoding="utf-8")
        (Path(d) / "queries.txt").write_text("Some text\\nSELECT id FROM orders;", encoding="utf-8")
        sub = Path(d) / "subdir"
        sub.mkdir()
        (sub / "nested.sql").write_text("UPDATE users SET active=1 WHERE id=5;", encoding="utf-8")
        yield d

class TestScanner:
    def test_finds_files(self, scanner, tmpdir_with_files):
        r = scanner.scan_directory(tmpdir_with_files)
        assert r.total_files_scanned > 0

    def test_recursive(self, scanner, tmpdir_with_files):
        r = scanner.scan_directory(tmpdir_with_files)
        assert "nested.sql" in [f.file_name for f in r.scanned_files]

    def test_extension_filter(self, scanner, tmpdir_with_files):
        r = scanner.scan_directory(tmpdir_with_files, extensions=[".sql"])
        assert all(f.extension == ".sql" for f in r.scanned_files)

    def test_nonexistent(self, scanner):
        r = scanner.scan_directory("/nonexistent/path/xyz")
        assert r.total_files_scanned == 0

    def test_content_read(self, scanner, tmpdir_with_files):
        r = scanner.scan_directory(tmpdir_with_files, extensions=[".sql"])
        assert all(len(f.content) > 0 for f in r.scanned_files)

    def test_keyword_filter(self, scanner, tmpdir_with_files):
        r = scanner.scan_directory(tmpdir_with_files, keyword_filter=["UPDATE"])
        for f in r.scanned_files:
            assert "UPDATE" in f.content.upper()

    def test_metadata_populated(self, scanner, tmpdir_with_files):
        r = scanner.scan_directory(tmpdir_with_files, extensions=[".sql"])
        for f in r.scanned_files:
            assert f.file_name and f.folder_path and f.size_bytes > 0

    def test_statistics(self, scanner, tmpdir_with_files):
        r = scanner.scan_directory(tmpdir_with_files)
        assert r.total_files_found >= r.total_files_scanned
'''

    # ================================================================
    # tests/test_formatter.py
    # ================================================================
    files["sql_consolidator/tests/test_formatter.py"] = '''"""Unit tests for SQL Formatter module."""
import pytest
import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from formatter.sql_formatter import SQLFormatter
from sql_parser.sql_parser import SQLParser, ExtractedQuery, QueryType


@pytest.fixture
def config():
    return {"formatter": {"keyword_case":"upper","indent_width":4,"reindent":True,"strip_comments":False}}

@pytest.fixture
def formatter(config):
    return SQLFormatter(config)

def make_q(sql, qtype=QueryType.SELECT):
    p = SQLParser({"parser": {}})
    n = p._normalize_sql(sql)
    fp = p._generate_fingerprint(n)
    q = ExtractedQuery(raw_sql=sql, normalized_sql=n, query_type=qtype,
                       source_file="/test/sample.sql", source_folder="/test",
                       line_number_start=1, line_number_end=5, fingerprint=fp,
                       complexity_score=3, importance_score=7, risk_level="LOW",
                       table_names=["users"])
    return q

class TestFormatter:
    def test_format_returns_content(self, formatter):
        q = make_q("select * from users where active=1")
        assert len(formatter.format_query(q)) > 0

    def test_ends_with_semicolon(self, formatter):
        q = make_q("SELECT * FROM users")
        assert formatter.format_query(q).rstrip().endswith(";")

    def test_metadata_comment(self, formatter):
        q = make_q("SELECT id FROM users")
        result = formatter._add_metadata_comment(q, "SELECT id FROM users;")
        assert "Source File" in result and "Risk Level" in result

    def test_select_category(self, formatter):
        assert "SELECT" in formatter._get_category_name(QueryType.SELECT)

    def test_ddl_category(self, formatter):
        assert "CREATE" in formatter._get_category_name(QueryType.CREATE) or "DDL" in formatter._get_category_name(QueryType.CREATE)

    def test_groups_by_type(self, formatter):
        queries = [make_q("SELECT * FROM users"), make_q("INSERT INTO t VALUES (1)", QueryType.INSERT), make_q("SELECT id FROM orders")]
        cat = formatter.format_all_queries(queries)
        assert isinstance(cat, dict) and len(cat) >= 1

    def test_master_doc_header(self, formatter):
        queries = [make_q("SELECT * FROM users")]
        meta = {"scan_datetime":"2024-01-01 00:00:00","input_directory":"/test",
                "total_files_scanned":1,"total_files_failed":0,
                "total_queries_found":1,"total_unique_queries":1,
                "total_duplicates_removed":0,"query_type_distribution":{"SELECT":1}}
        doc = formatter.generate_master_sql_document(queries, meta)
        assert "SQL QUERY CONSOLIDATION MASTER DOCUMENT" in doc
'''

    # ================================================================
    # sample_input/sample1.sql
    # ================================================================
    files["sql_consolidator/sample_input/sample1.sql"] = (
        "-- Sample SQL File 1: Production Queries\n\n"
        "SELECT\n"
        "    u.id,\n"
        "    u.username,\n"
        "    u.email,\n"
        "    COUNT(o.id) AS total_orders,\n"
        "    SUM(o.total_amount) AS lifetime_value\n"
        "FROM users u\n"
        "LEFT JOIN orders o ON u.id = o.user_id\n"
        "WHERE u.is_active = 1\n"
        "GROUP BY u.id, u.username, u.email\n"
        "HAVING COUNT(o.id) > 0\n"
        "ORDER BY lifetime_value DESC;\n\n"
        "-- DUPLICATE (different formatting - should be removed)\n"
        "select u.id, u.username, u.email,\n"
        "count(o.id) as total_orders, sum(o.total_amount) as lifetime_value\n"
        "from users u left join orders o on u.id = o.user_id\n"
        "where u.is_active = 1\n"
        "group by u.id, u.username, u.email\n"
        "having count(o.id) > 0 order by lifetime_value desc;\n\n"
        "INSERT INTO audit_log (user_id, action, timestamp)\n"
        "VALUES (@user_id, 'LOGIN', GETDATE());\n\n"
        "UPDATE users\n"
        "SET last_login = GETDATE(), login_count = login_count + 1\n"
        "WHERE id = @user_id;\n\n"
        "CREATE TABLE IF NOT EXISTS user_sessions (\n"
        "    session_id  VARCHAR(255) PRIMARY KEY,\n"
        "    user_id     INT NOT NULL,\n"
        "    created_at  TIMESTAMP DEFAULT CURRENT_TIMESTAMP,\n"
        "    expires_at  TIMESTAMP,\n"
        "    is_active   BOOLEAN DEFAULT TRUE,\n"
        "    FOREIGN KEY (user_id) REFERENCES users(id)\n"
        ");\n\n"
        "WITH monthly_revenue AS (\n"
        "    SELECT DATE_FORMAT(created_at, '%Y-%m') AS month,\n"
        "           SUM(amount) AS revenue, COUNT(*) AS order_count\n"
        "    FROM orders WHERE status = 'completed'\n"
        "    GROUP BY DATE_FORMAT(created_at, '%Y-%m')\n"
        ")\n"
        "SELECT * FROM monthly_revenue ORDER BY revenue DESC;\n\n"
        "-- RISKY: DELETE without WHERE\n"
        "DELETE FROM temp_processing_queue;\n\n"
        "DROP TABLE IF EXISTS old_backup_2019;\n\n"
        "ALTER TABLE users ADD COLUMN phone VARCHAR(20) NULL;\n"
    )

    # ================================================================
    # sample_input/sample2.txt
    # ================================================================
    files["sql_consolidator/sample_input/sample2.txt"] = (
        "Application Log File - 2024-01-15\n"
        "==================================\n"
        "[INFO] Starting batch process\n\n"
        "SELECT p.product_id, p.name, p.price,\n"
        "       c.category_name,\n"
        "       COALESCE(inv.quantity, 0) AS stock_quantity\n"
        "FROM products p\n"
        "INNER JOIN categories c ON p.category_id = c.id\n"
        "LEFT JOIN inventory inv ON p.product_id = inv.product_id\n"
        "WHERE p.is_active = 1 AND p.price > 0\n"
        "ORDER BY c.category_name, p.name;\n\n"
        "[INFO] 1542 records returned\n\n"
        "UPDATE inventory\n"
        "SET quantity = quantity - @sold_quantity,\n"
        "    last_updated = GETDATE()\n"
        "WHERE product_id = @product_id\n"
        "  AND quantity >= @sold_quantity;\n\n"
        "INSERT INTO order_items (order_id, product_id, quantity, unit_price)\n"
        "SELECT @order_id, product_id, @qty, price\n"
        "FROM products\n"
        "WHERE product_id = @product_id;\n\n"
        "MERGE INTO customer_summary AS target\n"
        "USING (\n"
        "    SELECT customer_id, SUM(amount) AS total_spent, COUNT(*) AS order_count\n"
        "    FROM orders\n"
        "    WHERE order_date >= DATEADD(MONTH, -1, GETDATE())\n"
        "    GROUP BY customer_id\n"
        ") AS source ON target.customer_id = source.customer_id\n"
        "WHEN MATCHED THEN\n"
        "    UPDATE SET target.total_spent = source.total_spent,\n"
        "               target.last_updated = GETDATE()\n"
        "WHEN NOT MATCHED THEN\n"
        "    INSERT (customer_id, total_spent, last_updated)\n"
        "    VALUES (source.customer_id, source.total_spent, GETDATE());\n\n"
        "[INFO] Batch process complete.\n"
    )

    # ================================================================
    # sample_input/sample3.log
    # ================================================================
    files["sql_consolidator/sample_input/sample3.log"] = (
        "2024-01-15 10:00:00 [DEBUG] Simple query executed:\n"
        "SELECT * FROM users\n\n"
        "2024-01-15 10:00:01 [INFO] Window function report:\n"
        "SELECT d.department_name, e.employee_id,\n"
        "       e.first_name + ' ' + e.last_name AS full_name,\n"
        "       e.salary,\n"
        "       AVG(e.salary) OVER (PARTITION BY d.department_id) AS dept_avg,\n"
        "       RANK() OVER (PARTITION BY d.department_id ORDER BY e.salary DESC) AS rank\n"
        "FROM employees e\n"
        "INNER JOIN departments d ON e.department_id = d.department_id\n"
        "WHERE e.hire_date >= '2020-01-01'\n"
        "ORDER BY d.department_name, rank;\n\n"
        "2024-01-15 10:00:05 [INFO] Transaction block:\n"
        "BEGIN TRANSACTION;\n"
        "    UPDATE accounts SET balance = balance - @amount WHERE account_id = @from;\n"
        "    UPDATE accounts SET balance = balance + @amount WHERE account_id = @to;\n"
        "    INSERT INTO transactions (from_account, to_account, amount, txn_date)\n"
        "    VALUES (@from, @to, @amount, GETDATE());\n"
        "COMMIT TRANSACTION;\n\n"
        "2024-01-15 10:00:10 [INFO] Stored function:\n"
        "CREATE OR REPLACE FUNCTION calculate_discount(p_customer_id INT, p_amount DECIMAL)\n"
        "RETURNS DECIMAL AS\n"
        "BEGIN\n"
        "    DECLARE @rate DECIMAL = 0.0;\n"
        "    DECLARE @cnt INT;\n"
        "    SELECT @cnt = COUNT(*) FROM orders\n"
        "    WHERE customer_id = p_customer_id AND status = 'completed';\n"
        "    IF @cnt >= 10 SET @rate = 0.15;\n"
        "    ELSE IF @cnt >= 5 SET @rate = 0.10;\n"
        "    RETURN p_amount * @rate;\n"
        "END;\n"
    )

    # ================================================================
    # run.bat
    # ================================================================
    files["sql_consolidator/run.bat"] = (
        "@echo off\n"
        "title SQL Query Consolidation Tool\n"
        "color 0A\n\n"
        "echo ============================================================\n"
        "echo   SQL QUERY CONSOLIDATION ^& DEDUPLICATION TOOL\n"
        "echo ============================================================\n"
        "echo.\n\n"
        "python --version >nul 2>&1\n"
        "if errorlevel 1 (\n"
        "    echo ERROR: Python is not installed or not in PATH\n"
        "    pause\n"
        "    exit /b 1\n"
        ")\n\n"
        "if not exist \"venv\" (\n"
        "    echo Creating virtual environment...\n"
        "    python -m venv venv\n"
        ")\n\n"
        "call venv\\Scripts\\activate.bat\n"
        "pip install -r requirements.txt -q\n"
        "if not exist \"logs\" mkdir logs\n"
        "if not exist \"output\" mkdir output\n\n"
        ":MENU\n"
        "echo.\n"
        "echo   1. Run CLI Tool\n"
        "echo   2. Launch Streamlit GUI\n"
        "echo   3. Run with Sample Input\n"
        "echo   4. Run Unit Tests\n"
        "echo   5. Exit\n"
        "echo.\n"
        "set /p choice=\"Enter choice (1-5): \"\n\n"
        "if \"%choice%\"==\"1\" goto CLI\n"
        "if \"%choice%\"==\"2\" goto GUI\n"
        "if \"%choice%\"==\"3\" goto SAMPLE\n"
        "if \"%choice%\"==\"4\" goto TEST\n"
        "if \"%choice%\"==\"5\" goto EXIT\n"
        "goto MENU\n\n"
        ":CLI\n"
        "set /p idir=\"Input directory: \"\n"
        "set /p odir=\"Output directory (default ./output): \"\n"
        "if \"%odir%\"==\"\" set odir=./output\n"
        "python main.py -i \"%idir%\" -o \"%odir%\" --verbose\n"
        "pause\n"
        "goto MENU\n\n"
        ":GUI\n"
        "echo Opening http://localhost:8501\n"
        "streamlit run ui/streamlit_app.py --server.port 8501\n"
        "goto MENU\n\n"
        ":SAMPLE\n"
        "python main.py -i \"./sample_input\" -o \"./output/sample_run\" --verbose\n"
        "echo Done! Check ./output/sample_run/\n"
        "pause\n"
        "goto MENU\n\n"
        ":TEST\n"
        "pytest tests/ -v --tb=short\n"
        "pause\n"
        "goto MENU\n\n"
        ":EXIT\n"
        "deactivate\n"
        "exit /b 0\n"
    )

    # ================================================================
    # run.sh
    # ================================================================
    files["sql_consolidator/run.sh"] = (
        "#!/bin/bash\n"
        "echo '============================================================'\n"
        "echo '  SQL QUERY CONSOLIDATION AND DEDUPLICATION TOOL'\n"
        "echo '============================================================'\n\n"
        "if ! command -v python3 &> /dev/null; then\n"
        "    echo 'ERROR: Python 3 not installed'; exit 1\n"
        "fi\n\n"
        "[ ! -d venv ] && python3 -m venv venv\n"
        "source venv/bin/activate\n"
        "pip install -r requirements.txt -q\n"
        "mkdir -p logs output\n\n"
        "echo '1. CLI Tool  2. Streamlit GUI  3. Sample Run  4. Tests'\n"
        "read -p 'Choice: ' c\n\n"
        "case $c in\n"
        "    1) read -p 'Input dir: ' i; read -p 'Output dir: ' o; o=${o:-./output}\n"
        "       python3 main.py -i \"$i\" -o \"$o\" --verbose ;;\n"
        "    2) streamlit run ui/streamlit_app.py ;;\n"
        "    3) python3 main.py -i ./sample_input -o ./output/sample_run --verbose ;;\n"
        "    4) pytest tests/ -v ;;\n"
        "    *) echo 'Invalid choice' ;;\n"
        "esac\n"
    )

    # ================================================================
    # README.md (safe version - no backtick code blocks)
    # ================================================================
    files["sql_consolidator/README.md"] = (
        "# SQL Query Consolidation and Deduplication Tool\n\n"
        "A professional Python application that scans multiple files containing SQL\n"
        "queries and generates a single consolidated document with only clean,\n"
        "unique, important queries.\n\n"
        "## Quick Start\n\n"
        "Step 1: Install dependencies\n"
        "    pip install -r requirements.txt\n\n"
        "Step 2: Run with sample data\n"
        "    python main.py -i ./sample_input -o ./output --verbose\n\n"
        "Step 3: Launch GUI\n"
        "    streamlit run ui/streamlit_app.py\n\n"
        "Step 4: Windows users can run run.bat\n\n"
        "## Command Line Options\n\n"
        "    python main.py -i INPUT_DIR -o OUTPUT_DIR [OPTIONS]\n\n"
        "    -i, --input        Input directory to scan (required)\n"
        "    -o, --output       Output directory for reports (required)\n"
        "    -e, --extensions   File extensions: .sql .txt .log\n"
        "    -k, --keywords     Filter keywords: production critical\n"
        "    --no-dedup         Skip deduplication\n"
        "    --min-complexity   Minimum complexity score (0-50)\n"
        "    --min-importance   Minimum importance score (0-30)\n"
        "    --verbose          Show detailed output\n\n"
        "## Output Files\n\n"
        "    master_consolidated_TIMESTAMP.sql   Main SQL document\n"
        "    summary_report_TIMESTAMP.txt        Statistics report\n"
        "    sql_report_TIMESTAMP.docx           Word document\n"
        "    sql_analysis_TIMESTAMP.xlsx         Excel workbook\n"
        "    duplicate_report_TIMESTAMP.txt      Duplicate details\n"
        "    error_log_TIMESTAMP.txt             Error log\n\n"
        "## Project Structure\n\n"
        "    sql_consolidator/\n"
        "    |-- main.py\n"
        "    |-- requirements.txt\n"
        "    |-- run.bat\n"
        "    |-- run.sh\n"
        "    |-- config/config.yaml\n"
        "    |-- scanner/file_scanner.py\n"
        "    |-- parser/sql_parser.py\n"
        "    |-- deduplicator/deduplicator.py\n"
        "    |-- formatter/sql_formatter.py\n"
        "    |-- analyzer/query_analyzer.py\n"
        "    |-- reports/report_generator.py\n"
        "    |-- ui/streamlit_app.py\n"
        "    |-- tests/\n"
        "    |-- sample_input/\n"
        "    |-- logs/\n"
        "    |-- output/\n\n"
        "## Running Tests\n\n"
        "    pytest tests/ -v\n"
        "    pytest tests/ --cov=. --cov-report=html\n\n"
        "## Requirements\n\n"
        "Python 3.8 or higher required.\n"
        "See requirements.txt for all dependencies.\n"
    )

    # ================================================================
    # Empty placeholder files
    # ================================================================
    files["sql_consolidator/logs/.gitkeep"] = ""
    files["sql_consolidator/output/.gitkeep"] = ""

    # ================================================================
    # Write ZIP
    # ================================================================
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for file_path, content in files.items():
            zf.writestr(file_path, content)

    zip_buffer.seek(0)
    output_zip = "sql_consolidator.zip"
    with open(output_zip, "wb") as f:
        f.write(zip_buffer.read())

    print("\n" + "=" * 60)
    print("  SQL Consolidator ZIP Generated Successfully!")
    print("=" * 60)
    print(f"  Output file : {output_zip}")
    print(f"  Total files : {len(files)}")
    print("\n  Files included:")
    for path in sorted(files.keys()):
        size = len(files[path].encode("utf-8"))
        print(f"    {path:<65} {size:>8} bytes")
    print("\n" + "=" * 60)
    print("  NEXT STEPS:")
    print("  1. Extract: sql_consolidator.zip")
    print("  2. Enter  : cd sql_consolidator")
    print("  3. Install: pip install -r requirements.txt")
    print("  4. Run    : python main.py -i ./sample_input -o ./output --verbose")
    print("  5. GUI    : streamlit run ui/streamlit_app.py")
    print("  6. Windows: run.bat")
    print("=" * 60 + "\n")


if __name__ == "__main__":
    create_zip()