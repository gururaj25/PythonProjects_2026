"""
Report Generator Module
Generates comprehensive reports in SQL, TXT, DOCX, and XLSX formats.
"""

import logging
from datetime import datetime
from typing import List, Dict
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

from formatter.sql_formatter import SQLFormatter

logger = logging.getLogger(__name__)


class ReportGenerator:
    def __init__(self, config: Dict, output_dir: str):
        self.config = config
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.formatter = SQLFormatter(config)
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    def generate_all_reports(
        self, queries: List, dedup_result, analysis_report, scan_metadata: Dict
    ) -> Dict[str, str]:
        output_files = {}
        cfg = self.config.get("output", {})

        if cfg.get("generate_sql", True):
            path = self._gen_master_sql(queries, scan_metadata)
            output_files["master_sql"] = path
            logger.info(f"Master SQL: {path}")

        if cfg.get("generate_txt", True):
            path = self._gen_text_report(queries, dedup_result, analysis_report, scan_metadata)
            output_files["text_report"] = path
            logger.info(f"Text report: {path}")

        if cfg.get("generate_docx", True):
            path = self._gen_word_doc(queries, dedup_result, analysis_report, scan_metadata)
            output_files["word_document"] = path
            logger.info(f"Word document: {path}")

        if cfg.get("generate_xlsx", True):
            path = self._gen_excel(queries, dedup_result, analysis_report, scan_metadata)
            output_files["excel_report"] = path
            logger.info(f"Excel report: {path}")

        output_files["duplicate_report"] = self._gen_dup_report(dedup_result)
        output_files["error_log"] = self._gen_error_log(scan_metadata)
        return output_files

    def _gen_master_sql(self, queries: List, metadata: Dict) -> str:
        path = str(self.output_dir / f"master_consolidated_{self.timestamp}.sql")
        content = self.formatter.generate_master_sql_document(queries, metadata)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return path

    def _gen_text_report(
        self, queries: List, dedup, analysis, metadata: Dict
    ) -> str:
        path = str(self.output_dir / f"summary_report_{self.timestamp}.txt")
        lines = [
            "=" * 80, "SQL QUERY CONSOLIDATION - SUMMARY REPORT", "=" * 80,
            f"Generated: {metadata.get('scan_datetime', 'N/A')}",
            "",
            "SCAN STATISTICS", "-" * 40,
            f"Input Directory    : {metadata.get('input_directory', 'N/A')}",
            f"Files Found        : {metadata.get('total_files_found', 0)}",
            f"Files Scanned      : {metadata.get('total_files_scanned', 0)}",
            f"Files Failed       : {metadata.get('total_files_failed', 0)}",
            "",
            "QUERY STATISTICS", "-" * 40,
            f"Total Queries Found    : {metadata.get('total_queries_found', 0)}",
            f"Unique Queries         : {dedup.total_unique}",
            f"Duplicates Removed     : {dedup.total_duplicates}",
            f"  Exact Duplicates     : {dedup.exact_duplicates}",
            f"  Semantic Duplicates  : {dedup.semantic_duplicates}",
            f"Deduplication Rate     : "
            f"{(dedup.total_duplicates / max(dedup.total_input, 1)) * 100:.1f}%",
            "",
            "QUERY TYPE DISTRIBUTION", "-" * 40,
        ]
        for qtype, count in analysis.query_type_distribution.items():
            lines.append(f"  {qtype:<30}: {count}")
        lines += ["", "RISK ANALYSIS", "-" * 40]
        for risk_level, count in analysis.risk_summary.items():
            lines.append(f"  {risk_level:<20}: {count}")
        lines += ["", "TOP REFERENCED TABLES", "-" * 40]
        for table, count in analysis.most_common_tables[:15]:
            lines.append(f"  {table:<40}: {count} references")
        lines += ["", analysis.natural_language_summary, "", "RECOMMENDATIONS", "-" * 40]
        for rec in analysis.recommendations:
            lines.append(f"  {rec}")
        if analysis.high_risk_queries:
            lines += ["", "HIGH RISK QUERIES", "-" * 40]
            for q in analysis.high_risk_queries[:10]:
                lines += [
                    f"  Risk: {q.risk_level} | Type: {q.query_type.value}",
                    f"  File: {q.source_file} (Line {q.line_number_start})",
                    f"  SQL : {q.raw_sql[:150]}",
                    "",
                ]
        lines.append("=" * 80)
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return path

    def _gen_word_doc(
        self, queries: List, dedup, analysis, metadata: Dict
    ) -> str:
        path = str(self.output_dir / f"sql_report_{self.timestamp}.docx")
        doc = Document()
        doc.core_properties.title = "SQL Query Consolidation Report"

        doc.add_heading("SQL Query Consolidation Report", 0)
        doc.add_paragraph(f"Generated: {metadata.get('scan_datetime', 'N/A')}")
        doc.add_paragraph(f"Input: {metadata.get('input_directory', 'N/A')}")

        doc.add_page_break()
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(
            f"Files scanned: {metadata.get('total_files_scanned', 0)} | "
            f"Queries found: {metadata.get('total_queries_found', 0)} | "
            f"Unique queries: {dedup.total_unique} | "
            f"Duplicates removed: {dedup.total_duplicates}"
        )

        doc.add_heading("Statistics", level=2)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.rows[0].cells[0].text = "Metric"
        tbl.rows[0].cells[1].text = "Value"
        stats = [
            ("Files Scanned", str(metadata.get("total_files_scanned", 0))),
            ("Total Queries Found", str(metadata.get("total_queries_found", 0))),
            ("Unique Queries", str(dedup.total_unique)),
            ("Duplicates Removed", str(dedup.total_duplicates)),
            ("Critical Risk", str(analysis.risk_summary.get("CRITICAL", 0))),
            ("High Risk", str(analysis.risk_summary.get("HIGH", 0))),
            ("Avg Complexity", f"{analysis.average_complexity:.1f}"),
        ]
        for metric, value in stats:
            row = tbl.add_row()
            row.cells[0].text = metric
            row.cells[1].text = value

        doc.add_page_break()
        doc.add_heading("Risk Analysis", level=1)
        if analysis.high_risk_queries:
            for q in analysis.high_risk_queries[:10]:
                doc.add_paragraph(
                    f"[{q.risk_level}] {q.query_type.value} | "
                    f"{Path(q.source_file).name} (Line {q.line_number_start}): "
                    f"{q.raw_sql[:100]}...",
                    style="List Bullet"
                )
        else:
            doc.add_paragraph("No high-risk queries detected.")

        doc.add_heading("Recommendations", level=1)
        for rec in analysis.recommendations:
            doc.add_paragraph(rec, style="List Bullet")

        doc.add_page_break()
        doc.add_heading("Consolidated SQL Queries", level=1)
        categorized = self.formatter.format_all_queries(queries)
        for category, qlist in sorted(categorized.items()):
            doc.add_heading(f"{category} ({len(qlist)})", level=2)
            for i, qsql in enumerate(qlist[:50], 1):
                p = doc.add_paragraph()
                run = p.add_run(f"-- Query {i}\n{qsql[:800]}")
                run.font.name = "Courier New"
                run.font.size = Pt(8)

        doc.save(path)
        return path

    def _gen_excel(
        self, queries: List, dedup, analysis, metadata: Dict
    ) -> str:
        path = str(self.output_dir / f"sql_analysis_{self.timestamp}.xlsx")
        wb = Workbook()

        # Summary sheet
        ws = wb.active
        ws.title = "Summary"
        ws["A1"] = "SQL QUERY CONSOLIDATION REPORT"
        ws["A1"].font = Font(bold=True, size=14, color="1F4E79")
        ws.merge_cells("A1:C1")
        ws["A2"] = f"Generated: {metadata.get('scan_datetime', 'N/A')}"
        ws.merge_cells("A2:C2")

        stats_rows = [
            ("", ""), ("SCAN STATISTICS", ""),
            ("Files Found", metadata.get("total_files_found", 0)),
            ("Files Scanned", metadata.get("total_files_scanned", 0)),
            ("Files Failed", metadata.get("total_files_failed", 0)),
            ("", ""), ("QUERY STATISTICS", ""),
            ("Total Queries Found", metadata.get("total_queries_found", 0)),
            ("Unique Queries", dedup.total_unique),
            ("Duplicates Removed", dedup.total_duplicates),
            ("Exact Duplicates", dedup.exact_duplicates),
            ("Semantic Duplicates", dedup.semantic_duplicates),
            ("", ""), ("RISK SUMMARY", ""),
            ("Critical Risk", analysis.risk_summary.get("CRITICAL", 0)),
            ("High Risk", analysis.risk_summary.get("HIGH", 0)),
            ("Medium Risk", analysis.risk_summary.get("MEDIUM", 0)),
            ("Low Risk", analysis.risk_summary.get("LOW", 0)),
            ("", ""), ("QUERY DISTRIBUTION", ""),
        ]
        for qtype, count in analysis.query_type_distribution.items():
            stats_rows.append((qtype, count))

        section_hdrs = {"SCAN STATISTICS", "QUERY STATISTICS", "RISK SUMMARY", "QUERY DISTRIBUTION"}
        row = 4
        for label, value in stats_rows:
            ws.cell(row=row, column=1, value=label)
            ws.cell(row=row, column=2, value=value)
            if label in section_hdrs:
                c = ws.cell(row=row, column=1)
                c.font = Font(bold=True, color="FFFFFF")
                c.fill = PatternFill(fill_type="solid", fgColor="2E75B6")
                ws.merge_cells(f"A{row}:B{row}")
            row += 1
        ws.column_dimensions["A"].width = 30
        ws.column_dimensions["B"].width = 20

        # All Queries sheet
        ws2 = wb.create_sheet("All Queries")
        headers = [
            "Query #", "Query Type", "Source File", "Folder",
            "Line Start", "Line End", "Tables", "Has JOINs",
            "Has Subquery", "Has CTE", "Complexity", "Importance",
            "Risk Level", "SQL Preview"
        ]
        hf = Font(bold=True, color="FFFFFF")
        hfill = PatternFill(fill_type="solid", fgColor="1F4E79")
        for col, hdr in enumerate(headers, 1):
            c = ws2.cell(row=1, column=col, value=hdr)
            c.font = hf
            c.fill = hfill
            c.alignment = Alignment(horizontal="center")
        ws2.freeze_panes = "A2"
        risk_colors = {
            "CRITICAL": "FF0000", "HIGH": "FFC000",
            "MEDIUM": "FFFF00", "LOW": "92D050"
        }
        for i, q in enumerate(queries, 1):
            r = i + 1
            vals = [
                i, q.query_type.value, q.source_file, q.source_folder,
                q.line_number_start, q.line_number_end,
                ", ".join(q.table_names[:5]),
                "Yes" if q.has_joins else "No",
                "Yes" if q.has_subquery else "No",
                "Yes" if q.has_cte else "No",
                q.complexity_score, q.importance_score, q.risk_level,
                q.raw_sql[:200].replace("\n", " "),
            ]
            for col, val in enumerate(vals, 1):
                cell = ws2.cell(row=r, column=col, value=val)
                if col == 13:
                    cell.fill = PatternFill(
                        fill_type="solid",
                        fgColor=risk_colors.get(str(val), "FFFFFF")
                    )
        ws2.auto_filter.ref = f"A1:{get_column_letter(len(headers))}1"
        for col, w in enumerate([8,20,50,40,10,10,40,10,12,10,12,12,12,60], 1):
            ws2.column_dimensions[get_column_letter(col)].width = w

        # High Risk sheet
        ws3 = wb.create_sheet("High Risk Queries")
        rh = ["Risk Level", "Query Type", "Source File", "Line", "SQL Preview", "Issue"]
        rf = PatternFill(fill_type="solid", fgColor="C00000")
        for col, h in enumerate(rh, 1):
            c = ws3.cell(row=1, column=col, value=h)
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = rf
        for i, q in enumerate(analysis.high_risk_queries, 2):
            issue = (
                "DELETE/UPDATE without WHERE - may affect ALL rows!"
                if q.risk_level == "CRITICAL"
                else "DROP/TRUNCATE - irreversible operation"
            )
            ws3.cell(row=i, column=1, value=q.risk_level)
            ws3.cell(row=i, column=2, value=q.query_type.value)
            ws3.cell(row=i, column=3, value=q.source_file)
            ws3.cell(row=i, column=4, value=q.line_number_start)
            ws3.cell(row=i, column=5, value=q.raw_sql[:200].replace("\n", " "))
            ws3.cell(row=i, column=6, value=issue)
        for col, w in enumerate([15, 20, 50, 10, 60, 50], 1):
            ws3.column_dimensions[get_column_letter(col)].width = w

        # Table Analysis sheet
        ws4 = wb.create_sheet("Table Analysis")
        th = ["Rank", "Table Name", "Reference Count"]
        tf = PatternFill(fill_type="solid", fgColor="375623")
        for col, h in enumerate(th, 1):
            c = ws4.cell(row=1, column=col, value=h)
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = tf
        for i, (tbl, cnt) in enumerate(analysis.most_common_tables, 2):
            ws4.cell(row=i, column=1, value=i - 1)
            ws4.cell(row=i, column=2, value=tbl)
            ws4.cell(row=i, column=3, value=cnt)

        # Duplicates sheet
        ws5 = wb.create_sheet("Duplicates")
        dh = ["Group #", "Dup Type", "Original File", "Orig Line",
              "Dup File", "Dup Line", "Match", "Preview"]
        df = PatternFill(fill_type="solid", fgColor="7030A0")
        for col, h in enumerate(dh, 1):
            c = ws5.cell(row=1, column=col, value=h)
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = df
        row = 2
        for i, grp in enumerate(dedup.duplicate_groups[:1000], 1):
            for dup in grp.get("duplicates", []):
                ws5.cell(row=row, column=1, value=i)
                ws5.cell(row=row, column=2, value=grp.get("duplicate_type", ""))
                ws5.cell(row=row, column=3, value=grp.get("original", ""))
                ws5.cell(row=row, column=4, value=grp.get("original_line", ""))
                ws5.cell(row=row, column=5, value=dup.get("file", ""))
                ws5.cell(row=row, column=6, value=dup.get("line", ""))
                ws5.cell(row=row, column=7, value=dup.get("type", ""))
                ws5.cell(row=row, column=8, value=grp.get("query_preview", "")[:150])
                row += 1

        wb.save(path)
        return path

    def _gen_dup_report(self, dedup) -> str:
        path = str(self.output_dir / f"duplicate_report_{self.timestamp}.txt")
        from deduplicator.deduplicator import QueryDeduplicator
        report = QueryDeduplicator(self.config).generate_dedup_report(dedup)
        with open(path, "w", encoding="utf-8") as f:
            f.write(report)
        return path

    def _gen_error_log(self, metadata: Dict) -> str:
        path = str(self.output_dir / f"error_log_{self.timestamp}.txt")
        lines = [f"ERROR LOG - {metadata.get('scan_datetime', 'N/A')}", "=" * 60, ""]
        failed = metadata.get("failed_files", [])
        if failed:
            lines.append(f"Failed Files ({len(failed)}):")
            for fi in failed:
                lines.append(f"  File : {fi.get('file_path', 'N/A')}")
                lines.append(f"  Error: {fi.get('error', 'Unknown')}")
                lines.append("")
        else:
            lines.append("No errors encountered.")
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return path
